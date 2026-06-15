"""第12G：报告模板解析服务（ReportTemplateService）。

职责：解析 SC/报告模版.docx，抽取「正式报告输出格式标准」：
- 9 章固定目录（标题以模板目录为准）；
- 四张三圈层量化表的指标行口径（核心/近邻/辐射）。

模板正文部分章节展开不完整（缺第6章案例参考、第9章附录），系统以「目录 9 章」为准，
自动补齐缺失章节。当模板文件缺失或解析失败时，回退到经核对确认的 canonical schema，
保证报告引擎始终可用。

红线：
- 仅本地读取 SC 模板；不外发、不入库；日志只输出结构统计，不输出模板原文。
- 模板只决定「结构与口径」，不决定任何事实数字；数字一律来自自研分析结果。
"""

from __future__ import annotations

import logging
from functools import lru_cache
from typing import Any

from app.config import settings

logger = logging.getLogger("cityrenew.report.template")

TEMPLATE_FILE = "报告模板.docx"

# 列口径（三圈层），全报告统一。
RING_COLUMNS = ["指标", "核心范围", "近邻范围", "辐射范围"]

# --------------------------------------------------------------------------- #
# Canonical schema（经核对 SC/报告模版.docx 确认；作为唯一标准与回退）
# --------------------------------------------------------------------------- #
CANONICAL_CHAPTERS: list[dict[str, str]] = [
    {"no": "1", "title": "项目基础概况", "note": "含城市更新适配性初步判断"},
    {"no": "2", "title": "区位与POI配套分析", "note": "依托POI兴趣点数据，结合更新需求"},
    {"no": "3", "title": "人口与客群画像分析", "note": "依托人口画像数据，匹配更新民生需求"},
    {"no": "4", "title": "房价与空间现状诊断", "note": "依托房价数据，识别更新潜力与痛点"},
    {"no": "5", "title": "产业与区域经济分析", "note": "依托产业数据，锚定更新功能定位"},
    {"no": "6", "title": "案例参考与政策适配分析", "note": "依托案例、政策数据，明确更新合规与路径"},
    {"no": "7", "title": "城市更新导向下的需求研判与潜力分析", "note": "依托全量数据"},
    {"no": "8", "title": "项目前策核心建议", "note": "含城市更新实施逻辑"},
    {"no": "9", "title": "附录：数据口径与样本说明", "note": ""},
]

# 四张量化表的行口径（与模板表头一致：核心/近邻/辐射）。
CANONICAL_TABLES: list[dict[str, Any]] = [
    {
        "key": "location",
        "title": "配套/区位指标表",
        "chapter_no": "2",
        "columns": RING_COLUMNS,
        "rows": [
            {"key": "transport", "label": "轨交/公交站点数量（个）"},
            {"key": "poi_total", "label": "POI兴趣点总数（个）"},
            {"key": "convenience", "label": "便民服务类POI（个）"},
            {"key": "commercial", "label": "商业服务类POI（个）"},
            {"key": "public", "label": "公共服务类POI（个）"},
        ],
    },
    {
        "key": "population",
        "title": "人口指标表",
        "chapter_no": "3",
        "columns": RING_COLUMNS,
        "rows": [
            {"key": "residential", "label": "居住人口（人）"},
            {"key": "worker", "label": "工作人口（人）"},
            {"key": "age_structure", "label": "年龄结构（0-18/19-35/36-59/60+占比）"},
            {"key": "consume_level", "label": "消费能力分级（占比）"},
            {"key": "income_level", "label": "收入水平分级（占比）"},
        ],
    },
    {
        "key": "housing",
        "title": "房价与空间现状表",
        "chapter_no": "4",
        "columns": RING_COLUMNS,
        "rows": [
            {"key": "secondhand_price", "label": "二手房均价（元/㎡）"},
            {"key": "secondhand_listing", "label": "二手房挂牌数（个）"},
            {"key": "rent_price", "label": "出租房均价（元/㎡）"},
            {"key": "rent_listing", "label": "出租房挂牌数（个）"},
            {"key": "building_age", "label": "建筑平均使用年限（年）"},
            {"key": "price_growth", "label": "房价历史涨幅（%）"},
        ],
    },
    {
        "key": "industry",
        "title": "产业/经济指标表",
        "chapter_no": "5",
        "columns": RING_COLUMNS,
        "rows": [
            {"key": "dominant_industry", "label": "主导产业类型"},
            {"key": "enterprise_count", "label": "产业企业数量（家）"},
            {"key": "economy_activity", "label": "区域经济活跃度指数"},
            {"key": "industry_pop_ratio", "label": "产业人口占比"},
        ],
    },
]


# --------------------------------------------------------------------------- #
# 模板解析（用于核对模板是否仍与 canonical 一致；不改变 canonical 口径）
# --------------------------------------------------------------------------- #
def _template_path():
    return settings.sc_path / TEMPLATE_FILE


def _parse_template_directory() -> dict[str, Any]:
    """读取模板 docx 的目录章节与表头，返回核对信息（不替换 canonical）。"""
    path = _template_path()
    info: dict[str, Any] = {"exists": path.exists(), "chapter_titles": [], "table_headers": []}
    if not path.exists():
        return info
    try:
        from docx import Document

        doc = Document(str(path))
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            # 目录形如 "1. xxx" ~ "9. xxx"
            if len(t) > 2 and t[0].isdigit() and t[1] in (".", "、", "．"):
                info["chapter_titles"].append(t)
        for tb in doc.tables:
            if tb.rows:
                header = " | ".join(c.text.strip() for c in tb.rows[0].cells)
                info["table_headers"].append(header)
    except Exception as exc:  # noqa: BLE001
        logger.warning("模板解析失败（回退 canonical）：%s", type(exc).__name__)
        info["error"] = type(exc).__name__
    return info


@lru_cache(maxsize=1)
def get_template_schema() -> dict[str, Any]:
    """返回报告的固定结构 schema（9 章 + 四张三圈层量化表）。

    始终以 canonical 为准（已对齐 SC/报告模版.docx）；附带模板核对信息，
    便于回归测试确认模板未发生结构性变化。
    """
    parsed = _parse_template_directory()
    schema = {
        "source": "SC/报告模板.docx",
        "template_present": parsed.get("exists", False),
        "chapters": CANONICAL_CHAPTERS,
        "tables": CANONICAL_TABLES,
        "ring_columns": RING_COLUMNS,
        "required_chapters": len(CANONICAL_CHAPTERS),
        "required_tables": len(CANONICAL_TABLES),
        "template_check": {
            "directory_chapters_found": len(parsed.get("chapter_titles", [])),
            "tables_found": len(parsed.get("table_headers", [])),
            "table_headers": parsed.get("table_headers", []),
        },
    }
    logger.info(
        "report template schema ready: present=%s dir_found=%s tables_found=%s",
        schema["template_present"],
        schema["template_check"]["directory_chapters_found"],
        schema["template_check"]["tables_found"],
    )
    return schema


def chapter_titles() -> list[str]:
    return [c["title"] for c in CANONICAL_CHAPTERS]
