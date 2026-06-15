"""报告导出（第7阶段，python-docx 确定性导出，绝不使用大模型/外部 API）。

把 report_content_service 生成的结构化报告内容导出为 docx，文件保存到
backend/data/outputs/reports/{project_id}/（已 gitignore，不入库、不外发）。

Word 中保留：章节标题、summary、关键要点、关键指标（含 evidence_id）、证据ID 列表、
数据局限。仅写入脱敏统计量与类别名，不含 raw_json / 原始点位 / 企业名 / 小区名 / 地址明细。
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document

from app.config import settings

logger = logging.getLogger("cityrenew.report.export")


def _report_dir(project_id: int) -> Path:
    d = settings.data_dir / "outputs" / "reports" / str(project_id)
    d.mkdir(parents=True, exist_ok=True)
    return d


def _safe_name(report_id: str) -> str:
    return report_id.replace(":", "_")


def _fmt_value(value: Any, unit: str | None) -> str:
    if value is None:
        return "数据不足"
    if isinstance(value, bool):
        text = str(value)
    elif isinstance(value, float):
        # 保真展示（不用 {:g}），与报告内容/落库口径一致。
        text = str(int(value)) if value == int(value) else f"{value:.4f}".rstrip("0").rstrip(".")
    else:
        text = str(value)
    return f"{text} {unit}".strip() if unit else text


def export_docx(content: dict[str, Any]) -> dict[str, Any]:
    """把结构化报告内容写为 docx，返回文件信息。"""
    project_id = content["project_id"]
    report_id = content["report_id"]

    doc = Document()
    doc.add_heading("城市更新前期策划大数据分析报告", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"项目：{content.get('project_name') or '—'}（ID {project_id}）\n").bold = True
    meta.add_run(f"项目类型：{content.get('project_type') or '数据不足'}\n")
    meta.add_run(f"报告编号：{report_id}\n")
    meta.add_run(f"生成时间(UTC)：{content.get('generated_at')}\n")
    meta.add_run(
        f"数据划分：{('/'.join(content.get('allowed_splits') or [])) or 'train/val'}"
        f"（used_test={content.get('used_test', False)}）"
    )

    for sec in content.get("sections", []):
        doc.add_heading(f"{sec.get('section_id')} {sec.get('title')}", level=1)

        if sec.get("summary"):
            doc.add_paragraph(sec["summary"])

        if sec.get("key_findings"):
            doc.add_heading("关键要点", level=2)
            for kf in sec["key_findings"]:
                doc.add_paragraph(kf, style="List Bullet")

        metrics = sec.get("metrics") or []
        if metrics:
            doc.add_heading("关键指标", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "指标"
            hdr[1].text = "数值"
            hdr[2].text = "证据ID"
            for m in metrics:
                row = table.add_row().cells
                row[0].text = str(m.get("label") or m.get("key") or "")
                row[1].text = _fmt_value(m.get("value"), m.get("unit"))
                row[2].text = str(m.get("evidence_id") or "")

        if sec.get("evidence_ids"):
            doc.add_heading("证据ID", level=2)
            doc.add_paragraph("；".join(sec["evidence_ids"]))

        if sec.get("data_limitations"):
            doc.add_heading("数据局限", level=2)
            for lim in sec["data_limitations"]:
                doc.add_paragraph(lim, style="List Bullet")

    if content.get("notes"):
        doc.add_heading("附：生成与合规说明", level=1)
        for n in content["notes"]:
            doc.add_paragraph(n, style="List Bullet")

    out_dir = _report_dir(project_id)
    file_name = _safe_name(report_id) + ".docx"
    file_path = out_dir / file_name
    doc.save(str(file_path))

    # 维护最新 docx 指针，便于 download 接口直接取用
    latest_path = out_dir / "latest.docx"
    doc.save(str(latest_path))

    size = file_path.stat().st_size if file_path.exists() else 0
    logger.info(
        "report exported project_id=%s report_id=%s file=%s size=%s",
        project_id, report_id, file_name, size,
    )
    return {
        "report_id": report_id,
        "project_id": project_id,
        "file_name": file_name,
        "file_path": str(file_path),
        "size_bytes": size,
        "download_url": f"/api/reports/{project_id}/download",
        "notes": [
            "报告文件已保存到 backend/data/outputs/reports/（已 gitignore，不入库、不外发）。",
            "导出为确定性渲染，未使用大模型、未调用外部 API。",
        ],
    }


def latest_docx_path(project_id: int) -> Path | None:
    """返回最新 docx 路径（无则 None）。"""
    path = settings.data_dir / "outputs" / "reports" / str(project_id) / "latest.docx"
    return path if path.exists() else None
