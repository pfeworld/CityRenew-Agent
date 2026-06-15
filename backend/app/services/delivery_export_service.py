"""第9阶段交付材料导出服务（export-delivery）。

调用 final_eval_service 取得最终自评结构，落盘为交付材料到
backend/data/outputs/final_eval/（已 gitignore）：
1. final_eval_summary.json
2. final_eval_summary.md
3. delivery_checklist.md
4. final_eval_summary.docx（可选，python-docx）

红线：
- 输出目录被 .gitignore 覆盖；不包含原始语料、不包含 raw_json、不包含 test 明细。
- 只包含汇总指标、门禁结果、结论、风险、使用说明。
- 落盘前对每个文件文本做泄露扫描，命中则中止导出（不写该文件）。
- 不调外部 API、不使用大模型。
"""

from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from app.config import settings
from app.services import final_eval_service as fes

logger = logging.getLogger("cityrenew.delivery_export")

FORBIDDEN_TOKENS = fes.FORBIDDEN_TOKENS


def _out_dir() -> Path:
    d = settings.data_dir / "outputs" / "final_eval"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _utcnow() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _scan(text: str) -> list[str]:
    return [tok for tok in FORBIDDEN_TOKENS if tok in text]


def _metric_line(m: Any) -> str:
    if isinstance(m, dict) and "metric_name" in m:
        return f"- **{m['metric_name']}**：{m['current_value']}（门槛 {m['threshold']}，状态 {m['status']}）"
    return f"- {m}"


def _render_summary_md(result: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# CityRenew Agent · 第9阶段最终自评汇总")
    lines.append("")
    lines.append(f"- 生成时间：{_utcnow()}")
    lines.append(f"- 运行模式：{result.get('mode')}")
    lines.append(f"- **overall_status**：{result.get('overall_status')}")
    lines.append(f"- **can_submit**：{result.get('can_submit')}")
    lines.append("")
    lines.append("## 一、三大核心指标")
    cm = result.get("core_metrics", {})
    for key in ("retrieval_accuracy", "report_completeness", "data_consistency"):
        if key in cm:
            lines.append(_metric_line(cm[key]))
    lines.append("")
    lines.append("## 二、扩展指标")
    em = result.get("extended_metrics", {})
    mt = result.get("model_test_metrics", {})
    lines.append(f"- **housing_test_mape**：{mt.get('test_mape')}（MAE={mt.get('test_mae')}，"
                 f"test_count={mt.get('test_count')}，target {mt.get('target')}，状态 {mt.get('status')}）")
    lines.append(f"- **test_used_for_training**：{mt.get('test_used_for_training')}")
    lines.append(f"- **evidence_coverage**：{em.get('evidence_coverage', {}).get('current_value')}")
    lines.append(f"- **number_traceability**：{em.get('number_traceability', {}).get('current_value')}")
    lines.append(f"- **mutation_tests_pass**：{em.get('mutation_tests_pass')}")
    lines.append("")
    lines.append("## 三、检索评测")
    rq = result.get("retrieval_quality_metrics", {})
    lines.append(f"- total={rq.get('total')} correct={rq.get('correct')} accuracy={rq.get('accuracy')}")
    lines.append(f"- coverage_by_source_type={rq.get('coverage_by_source_type')}")
    lines.append(f"- failed_cases 数量={len(rq.get('failed_cases', []))}")
    lines.append("")
    lines.append("## 四、test 隔离与合规")
    iso = result.get("test_isolation_check", {})
    lines.append(f"- test_isolation_check：{iso.get('status')}")
    lines.append(f"- {iso.get('statement')}")
    lines.append(f"- external_api_calls={result.get('external_api_calls')}，"
                 f"llm_used_for_scoring={result.get('llm_used_for_scoring')}")
    lines.append(f"- leakage_check：{result.get('leakage_check', {}).get('status')}")
    lines.append(f"- frontend_structure_check：{result.get('frontend_structure_check')}")
    lines.append(f"- frontend_demo_status_runtime：{result.get('frontend_demo_status_runtime')}")
    lines.append("")
    lines.append("## 五、风险")
    for r in result.get("risks", []) or ["（无）"]:
        lines.append(f"- {r}")
    lines.append("")
    lines.append("## 六、建议")
    for r in result.get("recommendations", []) or ["（无）"]:
        lines.append(f"- {r}")
    lines.append("")
    lines.append("## 七、结论")
    lines.append(result.get("final_summary", ""))
    lines.append("")
    lines.append(f"> test 使用策略：{result.get('test_policy')}")
    lines.append("")
    return "\n".join(lines)


def _render_checklist_md(result: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# 交付材料清单（delivery_checklist）")
    lines.append("")
    lines.append(f"- 生成时间：{_utcnow()}")
    lines.append(f"- delivery_checklist_complete（自动项）：{result.get('delivery_checklist_complete')}")
    lines.append("")
    lines.append("| 交付项 | 状态 | 是否需人工 | 说明 |")
    lines.append("|---|---|---|---|")
    for it in result.get("delivery_checklist", []):
        lines.append(f"| {it['item']} | {it['status']} | {it['requires_manual']} | {it['note']} |")
    lines.append("")
    lines.append("## 需人工提交项")
    for it in result.get("manual_pending_items", []) or ["（无）"]:
        lines.append(f"- {it}")
    lines.append("")
    lines.append("## 使用说明")
    lines.append("- 本目录材料由 final_eval_service 自动生成，仅含汇总指标/门禁结果/结论/风险/使用说明。")
    lines.append("- 不含原始语料、不含原始JSON、不含 test 明细；目录已被 .gitignore 覆盖，严禁提交。")
    lines.append("- 前端演示录屏/截图与 KupasEval 自评材料截图需人工补充后一并提交。")
    lines.append("")
    return "\n".join(lines)


def _write_docx(result: dict[str, Any], path: Path) -> bool:
    try:
        from docx import Document
    except Exception:  # noqa: BLE001
        return False
    doc = Document()
    doc.add_heading("CityRenew Agent · 第9阶段最终自评汇总", level=0)
    doc.add_paragraph(f"生成时间：{_utcnow()}")
    doc.add_paragraph(f"overall_status：{result.get('overall_status')}  |  can_submit：{result.get('can_submit')}")

    doc.add_heading("三大核心指标", level=1)
    cm = result.get("core_metrics", {})
    for key in ("retrieval_accuracy", "report_completeness", "data_consistency"):
        m = cm.get(key)
        if isinstance(m, dict):
            doc.add_paragraph(f"{m['metric_name']}：{m['current_value']}（门槛 {m['threshold']}，{m['status']}）",
                              style="List Bullet")

    doc.add_heading("扩展指标", level=1)
    mt = result.get("model_test_metrics", {})
    doc.add_paragraph(f"housing_test_mape：{mt.get('test_mape')}（MAE={mt.get('test_mae')}，"
                      f"test_count={mt.get('test_count')}，{mt.get('status')}）", style="List Bullet")
    doc.add_paragraph(f"test_used_for_training：{mt.get('test_used_for_training')}", style="List Bullet")

    doc.add_heading("test 隔离与合规", level=1)
    iso = result.get("test_isolation_check", {})
    doc.add_paragraph(iso.get("statement", ""))
    doc.add_paragraph(f"leakage_check：{result.get('leakage_check', {}).get('status')}；"
                      f"external_api_calls：{result.get('external_api_calls')}；"
                      f"llm_used_for_scoring：{result.get('llm_used_for_scoring')}")

    doc.add_heading("风险与建议", level=1)
    for r in result.get("risks", []):
        doc.add_paragraph(r, style="List Bullet")
    for r in result.get("recommendations", []):
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("结论", level=1)
    doc.add_paragraph(result.get("final_summary", ""))
    doc.add_paragraph(f"test 使用策略：{result.get('test_policy')}")
    doc.save(str(path))
    return True


def export_delivery(db: Session, include_docx: bool = True) -> dict[str, Any]:
    """生成最终自评 + 落盘交付材料到 outputs/final_eval。"""
    result = fes.run_final_eval(db)
    out = _out_dir()

    files_written: list[dict[str, Any]] = []
    skipped: list[dict[str, Any]] = []

    # 1. JSON
    json_text = json.dumps(result, ensure_ascii=False, indent=2, default=str)
    # 2. summary md / 3. checklist md
    summary_md = _render_summary_md(result)
    checklist_md = _render_checklist_md(result)

    text_files = {
        "final_eval_summary.json": json_text,
        "final_eval_summary.md": summary_md,
        "delivery_checklist.md": checklist_md,
    }

    for name, text in text_files.items():
        hits = _scan(text)
        if hits:
            skipped.append({"file": name, "reason": "leakage_detected", "hit_tokens": hits})
            continue
        path = out / name
        path.write_text(text, encoding="utf-8")
        files_written.append({"name": name, "bytes": path.stat().st_size})

    # 4. docx（可选）
    docx_ok = False
    if include_docx and not skipped:
        docx_path = out / "final_eval_summary.docx"
        docx_ok = _write_docx(result, docx_path)
        if docx_ok:
            files_written.append({"name": "final_eval_summary.docx", "bytes": docx_path.stat().st_size})

    export_success = bool(files_written) and not skipped

    logger.info(
        "delivery export: dir=%s written=%s skipped=%s docx=%s overall=%s can_submit=%s",
        out.name, len(files_written), len(skipped), docx_ok,
        result.get("overall_status"), result.get("can_submit"),
    )

    return {
        "mode": settings.app_mode,
        "phase": "9",
        "export_success": export_success,
        "output_dir": str(out.relative_to(settings.data_dir.parent)),
        "gitignore_covered": True,
        "files": files_written,
        "skipped_files": skipped,
        "leakage_check": result.get("leakage_check"),
        "overall_status": result.get("overall_status"),
        "can_submit": result.get("can_submit"),
        "notes": [
            "输出目录 backend/data/outputs/final_eval 已被 .gitignore 覆盖（backend/data/outputs/）。",
            "导出物仅含汇总指标/门禁结果/结论/风险/使用说明；不含原始语料、原始JSON、test 明细。",
            "落盘前对每个文本文件做泄露扫描，命中即跳过该文件并记入 skipped_files。",
        ],
    }
