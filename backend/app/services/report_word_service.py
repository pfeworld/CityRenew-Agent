"""基于 SC/报告模板.docx「就地填空」生成前策报告 Word，并由 LibreOffice 转 PDF（同源一致）。

设计原则（按用户返工要求）：
- 模板是「填空表单」：固定标题、说明文字、四张表、`____` 横线、`□` 选项一律原样保留，
  智能体只把横线 / 表格空格按本地分析结果填上，不重排、不增删模板既有结构；
- 封面：去掉 `#`、把「【XX市XX区XX项目】」替换为真实项目名、保留「数据来源：黑客松比赛提供专用数据库」；
- 「一、报告封面 / 三、报告正文」是模板结构说明，不作为正文显示（删除）；「二、报告目录」显示为「报告目录」；
- 模板正文只到第 7 节且缺案例章 / 附录：据目录补齐——在「需求研判」节前插入第 6 章「案例参考与政策适配分析」，
  末尾追加第 9 章「附录：数据口径与样本说明」，并把模板正文原 6/7 节顺延为 7/8 节，使目录与正文一致；
- 第 8、9 章无固定模板格式，按分析结果正常成文；
- 数字与结论来自 builder content（本地分析 / 模型 / 用户输入 / 附件），缺数据据实标注，绝不以 0 冒充；
- 生成前校验 model_run_id 与 analysis_result，否则 fail-closed；PDF 必须由 Word 经 soffice 转换，不用 ReportLab。
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

import docx
from docx.oxml.ns import qn
from docx.shared import Pt

from app.config import settings

logger = logging.getLogger("cityrenew.report.word")

TEMPLATE_NAME = "报告模板.docx"
GENERATOR_NAME = "CityRenew Agent 城市更新前期策划智能体"
TBD = "待补充"
_BLANK_RE = re.compile(r"_{2,}")
_SNAKE_RE = re.compile(r"[A-Za-z][A-Za-z0-9]*(?:_[A-Za-z0-9]+)+")

# 模板正文「需求研判」章主标题（其前插入第 6 章；其本身顺延为第 7 章）
_BODY_NEEDS_HEADING = "6. 城市更新导向下的需求研判与潜力分析"
# 章节 / 小节标题重编号映射（仅按完整文本匹配，避免误伤目录项）
_RENUMBER = {
    _BODY_NEEDS_HEADING: "7. 城市更新导向下的需求研判与潜力分析",
    "6.1 核心需求": "7.1 核心需求",
    "6.2 核心潜力": "7.2 核心潜力",
    "7. 项目前策核心建议": "8. 项目前策核心建议",
    "7.1 总体定位建议（含城市更新导向）": "8.1 总体定位建议（含城市更新导向）",
    "7.2 空间优化建议（城市更新改造重点）": "8.2 空间优化建议（城市更新改造重点）",
    "7.3 配套完善建议（民生导向）": "8.3 配套完善建议（民生导向）",
    "7.4 功能提升建议（产城融合）": "8.4 功能提升建议（产城融合）",
    "7.5 实施重点建议（城市更新路径）": "8.5 实施重点建议（城市更新路径）",
}
# 7.1-7.5（前策核心建议）横线按 guidance 关键词匹配填充
_ADVICE_KEYS = ["明确项目总体定位", "针对老旧建筑、低效空间", "结合城市更新民生需求",
                "结合产业适配性", "明确城市更新实施优先级"]


class ReportGateError(Exception):
    """生成前置校验失败（fail-closed）。"""


# --------------------------------------------------------------------------- #
# 路径
# --------------------------------------------------------------------------- #
def _template_path() -> Path:
    p = settings.sc_path / TEMPLATE_NAME
    if not p.exists():
        alt = settings.sc_path / "报告模版.docx"
        if alt.exists():
            return alt
    return p


def _report_dir(project_id: int) -> Path:
    d = settings.data_dir / "outputs" / "reports_v2" / str(project_id)
    d.mkdir(parents=True, exist_ok=True)
    return d


def safe_name(report_id: str) -> str:
    return report_id.replace(":", "_")


def file_for_report(report_id: str, ext: str) -> Path | None:
    try:
        pid = int(report_id.split(":p", 1)[1].split(":", 1)[0])
    except (IndexError, ValueError):
        return None
    path = _report_dir(pid) / f"{safe_name(report_id)}.{ext}"
    return path if path.exists() else None


# --------------------------------------------------------------------------- #
# 文本与段落处理
# --------------------------------------------------------------------------- #
def _clean(text: Any) -> str:
    s = str(text if text is not None else "").strip()
    s = _SNAKE_RE.sub("", s)
    s = s.replace("**", "").replace("*", "")
    s = re.sub(r"#{1,6}\s*", "", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


def _first_run_font(p):
    for r in p.runs:
        return (r.font.size, r.font.name, r.font.bold)
    return (None, None, None)


def _set_text(p, text: str) -> None:
    """重写整段文字，尽量保留原首个 run 的字体属性。"""
    size, name, bold = _first_run_font(p)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(_clean(text))
    if size:
        run.font.size = size
    if name:
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None:
        run.font.bold = bold


def _fill_one_blank(p, fill: str) -> None:
    """把段内第一处 `____` 替换为填充文本，其余文字（含括号说明）保持不变。"""
    _set_text(p, _BLANK_RE.sub(lambda m: _clean(fill), p.text, count=1))


def _delete_para(p) -> None:
    p._element.getparent().remove(p._element)


def _style_para(p, text: str, *, size: int, bold: bool, bullet: bool) -> None:
    pf = p.paragraph_format
    pf.space_after = Pt(3 if bullet else 6)
    pf.space_before = Pt(0)
    if bullet:
        pf.left_indent = Pt(12)
    run = p.add_run(("· " if bullet else "") + _clean(text))
    run.font.size = Pt(size)
    run.bold = bold


def _chapter_lines(c: dict, base_no: str) -> list[tuple[str, int, bool, bool]]:
    """把 builder 章节 dict 拍平成（文本, 字号, 加粗, 是否项目符号）。"""
    lines: list[tuple[str, int, bool, bool]] = [(f"{base_no}. {c['title']}", 15, True, False)]
    for para in c.get("paragraphs", []):
        lines.append((para, 11, False, False))
    for b in c.get("bullets", []):
        lines.append((b, 11, False, True))
    for sec in c.get("sections", []):
        lines.append((sec.get("title", ""), 13, True, False))
        for para in sec.get("paragraphs", []):
            lines.append((para, 11, False, False))
        for b in sec.get("bullets", []):
            lines.append((b, 11, False, True))
    return lines


def _insert_chapter_before(ref_para, c: dict, base_no: str) -> None:
    for text, size, bold, bullet in _chapter_lines(c, base_no):
        if not str(text).strip():
            continue
        np = ref_para.insert_paragraph_before("")
        _style_para(np, text, size=size, bold=bold, bullet=bullet)


def _append_chapter(doc, c: dict, base_no: str) -> None:
    doc.add_page_break()
    for text, size, bold, bullet in _chapter_lines(c, base_no):
        if not str(text).strip():
            continue
        _style_para(doc.add_paragraph(), text, size=size, bold=bold, bullet=bullet)


# --------------------------------------------------------------------------- #
# 表格填充（按模板行顺序，逐格写入，保留单位后缀，缺失据实标注）
# --------------------------------------------------------------------------- #
def _ordered_tables(chapters: list[dict]) -> list[dict]:
    out: list[dict] = []
    for c in chapters:
        out.extend(c.get("tables", []) or [])
    return out


def _set_cell(cell, value: str) -> None:
    orig = cell.text.strip()
    suffix = "年" if orig.endswith("年") else ("%" if orig.endswith("%") else "")
    v = _clean(value)
    if suffix and re.match(r"^\d", v):
        v = v + suffix
    size, name, _b = _first_run_font(cell.paragraphs[0]) if cell.paragraphs else (None, None, None)
    cell.text = ""
    run = cell.paragraphs[0].add_run(v)
    run.font.size = size or Pt(10)
    if name:
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _fill_tables(doc, tables: list[dict]) -> None:
    for ti, tb in enumerate(tables):
        if ti >= len(doc.tables):
            break
        dt = doc.tables[ti]
        rows = tb.get("rows") or []
        for ri, rd in enumerate(rows, start=1):
            if ri >= len(dt.rows):
                break
            cells = dt.rows[ri].cells
            for j, key in enumerate(("core", "nearby", "radiation"), start=1):
                if j < len(cells):
                    _set_cell(cells[j], rd.get(key, ""))


# --------------------------------------------------------------------------- #
# 段落「就地填空」总调度
# --------------------------------------------------------------------------- #
def _process_paragraphs(doc, fills: dict, gen_time: str, ch6: dict) -> None:
    location_text = fills.get("location_text") or ""
    conc = fills.get("conclusions") or {}
    advice = fills.get("advice") or {}
    type_choice = fills.get("type_choice") or "城市更新类"
    inserted_ch6 = False

    for p in list(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue

        if t in ("一、报告封面", "三、报告正文"):
            _delete_para(p)
        elif t == "二、报告目录":
            _set_text(p, "报告目录")
            p.paragraph_format.page_break_before = True   # 封面独立成页，目录另起一页
        elif t == "1. 项目基础概况":                       # 正文起始章（非目录项）
            p.paragraph_format.page_break_before = True   # 目录独立成页，正文另起一页
        elif "（含城市更新逻辑）" in t:               # 封面副标题
            _set_text(p, fills.get("subtitle") or t)
        elif t.startswith("生成模型"):
            _set_text(p, f"生成模型：{GENERATOR_NAME}")
        elif t.startswith("生成时间"):
            _set_text(p, f"生成时间：{gen_time}")
        elif t.startswith("数据来源") and "高德开放POI" not in t:
            # 保留模板必填口径，并据实补充全市 POI 来源（房价/人口/产业仍为比赛专用数据库）
            _set_text(p, "数据来源：黑客松比赛提供专用数据库；POI兴趣点补充自高德开放数据（GCJ02，覆盖上海全市）")
        elif t.startswith("- 项目位置"):
            _set_text(p, f"- 项目位置：{location_text}")
        elif t.startswith("- 项目类型"):
            _set_text(p, t.replace("□" + type_choice, "☑" + type_choice, 1))
        elif "自动提炼项目核心特征" in t:             # 1.2 适配性横线
            _fill_one_blank(p, fills.get("ch1_adapt") or "")
        elif t == _BODY_NEEDS_HEADING and not inserted_ch6:
            _insert_chapter_before(p, ch6, base_no="6")  # 先插入第 6 章
            _set_text(p, _RENUMBER[t])                   # 再把本节顺延为第 7 章
            inserted_ch6 = True
        elif t in _RENUMBER:                          # 其余章节 / 小节重编号
            _set_text(p, _RENUMBER[t])
        elif "：____" in t or ": ____" in t:           # 结论类横线（label：____（说明））
            label = t.split("：", 1)[0].split(". ", 1)[-1].strip()
            if label in conc:
                _fill_one_blank(p, conc[label])
        elif t.startswith("____"):                    # 前策核心建议 7.1-7.5 独立横线
            for key in _ADVICE_KEYS:
                if key in t and key in advice:
                    _fill_one_blank(p, advice[key])
                    break


# --------------------------------------------------------------------------- #
# 渲染后抽取纯文本 / 结构块（供预览与正文一致）
# --------------------------------------------------------------------------- #
def _heading_level(p) -> int:
    """按 Word 段落的加粗+字号识别标题层级，供前台按正式版排版。

    0=正文；1=封面主标题(≥20)；2=「报告目录」与章标题(≥15)；3=小节标题(≥12.5)。
    数字编号的正文条目为非加粗11号字，归为正文，不误判为标题。
    """
    runs = p.runs
    if not any(r.bold for r in runs):
        return 0
    size = next((r.font.size.pt for r in runs if r.font and r.font.size), None)
    if size is None:
        return 0
    if size >= 20:
        return 1
    if size >= 15:
        return 2
    if size >= 12.5:
        return 3
    return 0


def _extract_blocks(doc) -> tuple[str, list[dict]]:
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}
    blocks: list[dict] = []
    lines: list[str] = []
    for child in body.iterchildren():
        if child in para_map:
            para = para_map[child]
            txt = para.text.strip()
            if txt:
                blocks.append({"type": "para", "text": txt, "level": _heading_level(para)})
                lines.append(txt)
        elif child in table_map:
            t = table_map[child]
            rows = [[c.text.strip() for c in r.cells] for r in t.rows]
            blocks.append({"type": "table", "rows": rows})
            for r in rows:
                lines.append(" | ".join(r))
    return "\n".join(lines), blocks


# --------------------------------------------------------------------------- #
# Word → PDF（LibreOffice，绝不使用 ReportLab）
# --------------------------------------------------------------------------- #
def _soffice_bin() -> str:
    for cand in ("soffice", "/opt/homebrew/bin/soffice",
                 "/Applications/LibreOffice.app/Contents/MacOS/soffice"):
        if shutil.which(cand) or Path(cand).exists():
            return cand
    return "soffice"


def _convert_to_pdf(docx_path: Path) -> Path:
    out_dir = docx_path.parent
    with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile:
        cmd = [
            _soffice_bin(), "--headless", "--norestore", "--convert-to", "pdf",
            "--outdir", str(out_dir),
            f"-env:UserInstallation=file://{profile}",
            str(docx_path),
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
    produced = out_dir / (docx_path.stem + ".pdf")
    if not produced.exists() or produced.stat().st_size == 0:
        logger.warning("soffice 转换失败 rc=%s err=%s", proc.returncode, (proc.stderr or "")[:200])
        raise ReportGateError("Word 转 PDF 失败，请确认 LibreOffice 可用。")
    return produced


# --------------------------------------------------------------------------- #
# 主入口
# --------------------------------------------------------------------------- #
def build_and_convert(content: dict, bundle: dict) -> dict[str, Any]:
    """以模板为基底「就地填空」生成 Word → 转 PDF。

    content：report_builder_service.build_report 产出（含 chapters / template_fills / report_id）。
    bundle ：model_inference_service.run_inference 产出（须含 model_run_id + analysis_result）。
    """
    if not bundle or bundle.get("status") != "ok":
        raise ReportGateError("尚未获得有效的本地分析结果，无法生成报告。")
    if not bundle.get("model_run_id"):
        raise ReportGateError("缺少模型运行编号，无法生成报告。")
    if not bundle.get("analysis_result"):
        raise ReportGateError("缺少结构化分析结果，无法生成报告。")

    chapters = content.get("chapters") or []
    fills = content.get("template_fills")
    if len(chapters) < 9 or not fills:
        raise ReportGateError("报告内容不完整，无法生成报告。")

    template = _template_path()
    if not template.exists():
        raise ReportGateError("未找到报告模板文件。")

    report_id = content["report_id"]
    pid = content["project_id"]
    out_dir = _report_dir(pid)
    docx_path = out_dir / f"{safe_name(report_id)}.docx"

    # 1) 复制模板作为基底（绝不覆盖源文件）
    shutil.copyfile(template, docx_path)
    doc = docx.Document(str(docx_path))

    # 2) 表格 + 段落「就地填空」；插入第 6 章、追加第 9 章
    gen_time = datetime.now().strftime("%Y年%m月%d日")
    _fill_tables(doc, _ordered_tables(chapters))
    _process_paragraphs(doc, fills, gen_time, chapters[5])   # chapters[5] = 案例参考与政策适配分析
    _append_chapter(doc, chapters[8], base_no="9")           # chapters[8] = 附录
    doc.save(str(docx_path))

    # 3) Word → PDF（同源）
    pdf_path = _convert_to_pdf(docx_path)

    # 4) 抽取渲染后文本 / 结构块，回写 content（保证预览与正文一致）
    rendered_text, rendered_blocks = _extract_blocks(docx.Document(str(docx_path)))
    content["rendered_text"] = rendered_text
    content["rendered_blocks"] = rendered_blocks
    try:
        from app.services import report_builder_service as rb
        rb._persist(content)
    except Exception:  # noqa: BLE001
        logger.warning("回写渲染文本失败（不影响生成）", exc_info=True)

    stats = {"chapters": 9, "tables": len(_ordered_tables(chapters))}
    logger.info("report word+pdf built report_id=%s docx=%s pdf=%s stats=%s run_id=%s",
                report_id, docx_path.name, pdf_path.name, stats, bundle.get("model_run_id"))
    return {
        "report_id": report_id,
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "docx_size": docx_path.stat().st_size,
        "pdf_size": pdf_path.stat().st_size,
        "pdf_from_word": True,
        "model_run_id": bundle.get("model_run_id"),
        "fill_stats": stats,
        "template_used": template.name,
    }
