"""文档解析服务（第3阶段：RAG 知识库）。

职责：把参考资料中的政策 / 模板 / 脱敏报告 / 字段口径说明等文档解析为
结构化知识块（ParsedChunk），供 rag_service 入库与建索引。

支持格式：docx / pdf / xlsx / txt / md。

红线：
- 仅本地读取；不调用任何外部 API。
- 日志只输出文件名、块数、类型分布，绝不输出语料原文。
- chunk_text 可本地入库；但摘要/关键词由本地规则生成（非 LLM），
  接口默认只暴露摘要与限长片段。
- 仅解析"允许进入知识库"的文档；本阶段不处理 PPT / 图片 / 仅参考类文档。
"""

from __future__ import annotations

import hashlib
import logging
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.config import settings

logger = logging.getLogger("cityrenew.parser")

# --------------------------------------------------------------------------- #
# 知识源清单（本阶段允许入库的文档）
# 每项：file_name, source_type, data_type(用于 evidence_id), split
# 本阶段文档型知识源统一标 train（不在 split_manifest 中，且不引入 test）。
# --------------------------------------------------------------------------- #
KNOWLEDGE_SOURCES: list[dict[str, str]] = [
    {
        "file_name": "上海市城市更新条例.pdf",
        "source_type": "policy",
        "data_type": "policy",
        "split": "train",
    },
    {
        "file_name": "报告模板.docx",
        "source_type": "template",
        "data_type": "template",
        "split": "train",
    },
    {
        "file_name": "XX项目大数据分析报告（脱敏版）.docx",
        "source_type": "case_report",
        "data_type": "case",
        "split": "train",
    },
    {
        "file_name": "矢量数据样例及说明表.xlsx",
        "source_type": "field_spec",
        "data_type": "spec",
        "split": "train",
    },
    {
        "file_name": "城市更新前期策划智能体训练及测试数据集说明.docx",
        "source_type": "dataset_spec",
        "data_type": "spec",
        "split": "train",
    },
]


@dataclass
class ParsedChunk:
    source_file: str
    source_type: str
    data_type: str
    section: str | None
    page_no: int | None
    chunk_text: str
    chunk_summary: str
    keywords: list[str]
    metadata: dict[str, Any] = field(default_factory=dict)
    is_sensitive: bool = True
    split: str = "train"
    chunk_id: str = ""
    evidence_id: str = ""


# --------------------------------------------------------------------------- #
# 文本工具（本地规则，非 LLM）
# --------------------------------------------------------------------------- #
_WS_RE = re.compile(r"[ \t\u3000]+")
_SENT_SPLIT_RE = re.compile(r"(?<=[。！？!?；;\n])")


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_RE.sub(" ", text)
    # 合并多余空行
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def summarize(text: str, max_chars: int | None = None) -> str:
    """本地规则摘要：取前若干句并截断，绝不调用 LLM。"""
    max_chars = max_chars or settings.rag_summary_max_chars
    text = _clean(text).replace("\n", " ")
    if len(text) <= max_chars:
        return text
    # 句子边界尽量优雅截断
    out = ""
    for sent in _SENT_SPLIT_RE.split(text):
        if not sent:
            continue
        if len(out) + len(sent) > max_chars:
            break
        out += sent
    if not out:
        out = text[:max_chars]
    return out.strip().rstrip("，,") + "…"


_STOPWORDS = {
    "的", "了", "和", "与", "及", "或", "在", "是", "为", "对", "等", "可", "并",
    "以", "其", "该", "本", "中", "上", "下", "之", "也", "于", "由", "向", "从",
    "应", "需", "如", "按", "将", "被", "把", "有", "无", "不", "我们", "进行",
    "通过", "包括", "以及", "根据", "其中", "相关", "情况", "方面", "一个",
}


def extract_keywords(text: str, top_k: int = 8) -> list[str]:
    """本地关键词抽取（jieba + 频次），失败时回退正则切词。"""
    text = _clean(text)
    tokens: list[str] = []
    try:
        import jieba  # 本地分词，无网络

        tokens = [t.strip() for t in jieba.cut(text) if len(t.strip()) > 1]
    except Exception:  # noqa: BLE001  jieba 不可用时回退
        tokens = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,}", text)

    counts: Counter[str] = Counter()
    for t in tokens:
        if t in _STOPWORDS:
            continue
        if t.isdigit():
            continue
        counts[t] += 1
    return [w for w, _ in counts.most_common(top_k)]


def _split_long(text: str) -> list[str]:
    """长文本按目标字数 + 重叠二次切块。"""
    text = _clean(text)
    max_chars = settings.rag_chunk_max_chars
    overlap = settings.rag_chunk_overlap_chars
    if len(text) <= max_chars:
        return [text] if text else []

    sents = [s for s in _SENT_SPLIT_RE.split(text) if s and s.strip()]
    chunks: list[str] = []
    buf = ""
    for sent in sents:
        if len(buf) + len(sent) > max_chars and buf:
            chunks.append(buf.strip())
            # 重叠：保留尾部 overlap 字符
            buf = buf[-overlap:] if overlap else ""
        buf += sent
    if buf.strip():
        chunks.append(buf.strip())
    return chunks


def _source_hash8(source_file: str) -> str:
    return hashlib.sha256(source_file.encode("utf-8")).hexdigest()[:8]


def _make_chunk_id(stem: str, idx: int) -> str:
    return f"{stem}-{idx:04d}"


# --------------------------------------------------------------------------- #
# 各格式解析器：返回 (section, page_no, raw_text) 列表
# --------------------------------------------------------------------------- #
def _parse_docx(path: Path) -> list[tuple[str | None, int | None, str]]:
    from docx import Document

    doc = Document(str(path))
    blocks: list[tuple[str | None, int | None, str]] = []
    current_section: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            text = "\n".join(buffer).strip()
            if text:
                blocks.append((current_section, None, text))
            buffer.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        is_heading = style.startswith("heading") or style.startswith("title")
        if is_heading:
            flush()
            current_section = text
            blocks.append((current_section, None, text))
        else:
            buffer.append(text)
    flush()

    # 表格内容（按行拼接，作为口径/数据说明）
    for ti, table in enumerate(doc.tables):
        rows_text: list[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            blocks.append((f"表格{ti + 1}", None, "\n".join(rows_text)))
    return blocks


def _parse_pdf(path: Path) -> list[tuple[str | None, int | None, str]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    blocks: list[tuple[str | None, int | None, str]] = []
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001  个别页解析失败不阻断
            text = ""
        text = _clean(text)
        if text:
            blocks.append((None, page_no, text))
    return blocks


def _parse_xlsx(path: Path) -> list[tuple[str | None, int | None, str]]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    blocks: list[tuple[str | None, int | None, str]] = []
    for ws in wb.worksheets:
        header: list[str] = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if not values:
                continue
            if not header:
                header = values
                blocks.append((ws.title, row_idx, " | ".join(values)))
                continue
            blocks.append((ws.title, row_idx, " | ".join(values)))
    wb.close()
    return blocks


def _parse_text(path: Path) -> list[tuple[str | None, int | None, str]]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [(None, None, text)] if text.strip() else []


_PARSERS = {
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
    ".xlsx": _parse_xlsx,
    ".txt": _parse_text,
    ".md": _parse_text,
}


# --------------------------------------------------------------------------- #
# 入口
# --------------------------------------------------------------------------- #
def parse_file(
    path: Path,
    source_type: str,
    data_type: str,
    split: str = "train",
) -> list[ParsedChunk]:
    """解析单个文档为 ParsedChunk 列表。"""
    ext = path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"不支持的文件类型: {ext}")

    source_file = path.name
    stem = path.stem
    shash = _source_hash8(source_file)
    raw_blocks = parser(path)

    chunks: list[ParsedChunk] = []
    idx = 0
    for section, page_no, raw_text in raw_blocks:
        for piece in _split_long(raw_text):
            if not piece or len(piece) < 2:
                continue
            chunk_id = _make_chunk_id(stem, idx)
            evidence_id = f"{data_type}:{shash}:{chunk_id}"
            chunks.append(
                ParsedChunk(
                    source_file=source_file,
                    source_type=source_type,
                    data_type=data_type,
                    section=section,
                    page_no=page_no,
                    chunk_text=piece,
                    chunk_summary=summarize(piece),
                    keywords=extract_keywords(piece),
                    metadata={"ext": ext, "char_len": len(piece)},
                    is_sensitive=True,
                    split=split,
                    chunk_id=chunk_id,
                    evidence_id=evidence_id,
                )
            )
            idx += 1

    logger.info("parsed %s: chunks=%s type=%s", source_file, len(chunks), source_type)
    return chunks


def parse_knowledge_sources() -> tuple[list[ParsedChunk], list[dict[str, Any]]]:
    """解析全部已配置知识源，返回 (chunks, 文件级报告)。

    报告仅含统计量（文件名/块数/类型/是否存在），不含原文。
    """
    ref_dir = settings.reference_path
    all_chunks: list[ParsedChunk] = []
    reports: list[dict[str, Any]] = []

    for src in KNOWLEDGE_SOURCES:
        path = ref_dir / src["file_name"]
        if not path.exists():
            reports.append(
                {
                    "file_name": src["file_name"],
                    "source_type": src["source_type"],
                    "exists": False,
                    "chunks": 0,
                    "note": "文件不存在，已跳过",
                }
            )
            continue
        try:
            chunks = parse_file(
                path, src["source_type"], src["data_type"], src["split"]
            )
        except Exception as exc:  # noqa: BLE001  单文件失败不阻断整体
            logger.warning("parse failed %s: %s", src["file_name"], type(exc).__name__)
            reports.append(
                {
                    "file_name": src["file_name"],
                    "source_type": src["source_type"],
                    "exists": True,
                    "chunks": 0,
                    "note": f"解析失败: {type(exc).__name__}",
                }
            )
            continue
        all_chunks.extend(chunks)
        reports.append(
            {
                "file_name": src["file_name"],
                "source_type": src["source_type"],
                "exists": True,
                "chunks": len(chunks),
                "split": src["split"],
            }
        )
    return all_chunks, reports
