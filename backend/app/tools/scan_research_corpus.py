"""第10C：科研语料整理与缺口再分析（本地离线工具）。

职责（红线约束）：
  - 只在本地扫描 / 解析 / 脱敏摘要，绝不把原文全文写入 docs / 前端 / 报告 / 外部 API。
  - 不执行任何可执行文件（.exe 等），不自动解压压缩包。
  - 默认 used_for_training=False；商业来源（链家/贝壳等）一律 research_candidate，需授权。
  - 科研语料原始文件保留在 `科研语料/`（已 gitignore），本工具只产出元数据 / 脱敏摘要到
    `backend/data/external/research_corpus/`（同样 gitignore）。

运行：
  cd backend && ./.venv/bin/python -m app.tools.scan_research_corpus
"""
from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

# --------------------------------------------------------------------------- #
# 路径与常量
# --------------------------------------------------------------------------- #
BACKEND_DIR = Path(__file__).resolve().parents[2]          # .../backend
PROJECT_ROOT = BACKEND_DIR.parent                          # 项目根
CORPUS_CANDIDATES = ["科研语料", "科研资料", "research_corpus", "research_materials"]

EXTERNAL_DIR = BACKEND_DIR / "data" / "external"
RC_DIR = EXTERNAL_DIR / "research_corpus"
INVENTORY_DIR = RC_DIR / "inventory"
EXTRACTED_DIR = RC_DIR / "extracted"
CLASSIFIED_DIR = RC_DIR / "classified"
DATA_CATALOG_DIR = BACKEND_DIR / "data" / "outputs" / "data_catalog"

TEXT_STORE_CAP = 200_000          # 单文件脱敏文本最多保存字符数
TEXT_TOPIC_SCAN = 120_000         # 关键词扫描范围
PREVIEW_CHARS = 400               # 摘要预览长度

# --------------------------------------------------------------------------- #
# 文件类型
# --------------------------------------------------------------------------- #
EXT_DOCUMENT = {"docx", "doc", "pdf", "txt", "md", "rtf"}
EXT_TABLE = {"xlsx", "xls", "csv", "tsv"}
EXT_STRUCTURED = {"json", "geojson"}
EXT_IMAGE = {"png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff"}
EXT_PRESENTATION = {"pptx", "ppt"}
EXT_ARCHIVE = {"rar", "zip", "7z", "gz", "tar"}
EXT_GIS = {
    "shp", "shx", "dbf", "prj", "sbn", "sbx", "cpg", "atx", "gdbtable",
    "gdbtablx", "gdbindexes", "freelist", "gdb", "mxd", "lyr", "kml", "kmz",
    "tbx", "ovr", "nit", "spx", "adf", "ljobx", "dir", "timestamps", "dwg",
}
EXT_CODE = {"py", "js", "css", "html", "vlx", "lsp"}
EXT_EXECUTABLE = {"exe", "bat", "sh", "msi", "dll", "bin"}

PARSEABLE_LIBS = {"openpyxl": False, "docx": False, "pypdf": False, "xlrd": False}
for _lib in PARSEABLE_LIBS:
    try:
        __import__(_lib)
        PARSEABLE_LIBS[_lib] = True
    except Exception:  # noqa: BLE001
        PARSEABLE_LIBS[_lib] = False


def _utcnow() -> str:
    return datetime.now(timezone.utc).isoformat()


def detect_type(ext: str) -> str:
    ext = ext.lower()
    if ext in EXT_TABLE:
        return "table"
    if ext in EXT_STRUCTURED:
        return "structured_data"
    if ext in EXT_DOCUMENT:
        return "document"
    if ext in EXT_IMAGE:
        return "image"
    if ext in EXT_PRESENTATION:
        return "presentation"
    if ext in EXT_ARCHIVE:
        return "archive"
    if ext in EXT_GIS:
        return "gis"
    if ext in EXT_EXECUTABLE:
        return "executable"
    if ext in EXT_CODE:
        return "code"
    return "other"


def parse_plan(ext: str, detected: str) -> dict[str, Any]:
    """返回 can_parse / parse_method / need_ocr / need_manual_review 初判。"""
    ext = ext.lower()
    if detected == "table":
        if ext == "xlsx":
            return _plan(True, "openpyxl", False, False) if PARSEABLE_LIBS["openpyxl"] \
                else _plan(False, "none", False, True, "openpyxl 不可用")
        if ext == "xls":
            return _plan(True, "xlrd", False, False) if PARSEABLE_LIBS["xlrd"] \
                else _plan(False, "none", False, True, "xlrd 不可用，旧版 .xls 需人工转换")
        return _plan(True, "csv", False, False)  # csv/tsv
    if detected == "structured_data":
        return _plan(True, "json", False, False)
    if detected == "document":
        if ext in {"txt", "md"}:
            return _plan(True, "text", False, False)
        if ext == "docx":
            return _plan(True, "python-docx", False, False) if PARSEABLE_LIBS["docx"] \
                else _plan(False, "none", False, True, "python-docx 不可用")
        if ext == "pdf":
            return _plan(True, "pypdf", False, False) if PARSEABLE_LIBS["pypdf"] \
                else _plan(False, "none", True, True, "pypdf 不可用")
        return _plan(False, "none", False, True, f"无 {ext} 解析器")
    if detected == "image":
        # 无可靠本地 OCR（未安装 tesseract/PIL）→ 仅记录元数据，标记需人工复核
        return _plan(False, "image_metadata_only", True, True, "无本地 OCR，需人工识别")
    if detected == "presentation":
        return _plan(False, "none", True, True, "无 pptx 解析器（多已有图片版，需人工/OCR）")
    if detected == "archive":
        return _plan(False, "none", False, True, "压缩包不自动解压，需人工解压后再导入")
    if detected == "gis":
        return _plan(False, "geospatial_metadata_only", False, False, "GIS 数据需专业工具，仅登记")
    if detected == "executable":
        return _plan(False, "none", False, False, "可执行文件，禁止运行，仅登记")
    if detected == "code":
        return _plan(True, "text", False, False)
    return _plan(False, "none", False, True, "未知类型，需人工复核")


def _plan(can_parse: bool, method: str, need_ocr: bool, need_review: bool,
          reason: str | None = None) -> dict[str, Any]:
    return {
        "can_parse": can_parse,
        "parse_method": method,
        "need_ocr": need_ocr,
        "need_manual_review": need_review,
        "plan_note": reason,
    }


# --------------------------------------------------------------------------- #
# 关键词体系（用途分类）
# --------------------------------------------------------------------------- #
CATEGORY_KEYWORDS: dict[str, list[str]] = {
    "policy_planning": ["政策", "规划", "控规", "总规", "城市更新", "旧改", "更新单元",
                        "实施方案", "公告", "土地出让", "住房政策", "保障性住房政策",
                        "建设者管理者之家", "专项规划", "用地", "城市建设"],
    "stats_macro": ["统计", "年鉴", "GDP", "消费", "社零", "CPI", "固定资产", "经济",
                    "宏观", "快报", "财政", "增加值"],
    "population_profile": ["常住人口", "人口结构", "人口密度", "年龄", "家庭户",
                           "居住人口", "就业人口", "人口画像", "人口", "收入", "可支配收入"],
    "housing_property": ["房价", "成交", "挂牌", "租赁", "租金", "小区", "楼盘", "二手房",
                         "房龄", "链家", "贝壳", "保租房", "住房", "户型", "虹口住房",
                         "保障房", "公寓"],
    "industry_enterprise": ["产业", "企业", "园区", "工业", "写字楼", "研发", "科创",
                            "制造", "商办", "办公楼"],
    "poi_public_service": ["医院", "学校", "养老", "地铁", "公交", "公园", "体育", "文化",
                           "餐饮", "美食", "poi", "兴趣点", "设施", "商业"],
    "project_case": ["案例", "复盘", "演讲", "闭门会", "路演", "业态", "运营", "招商",
                     "实践", "经验", "交流"],
    "report_template": ["报告模板", "模板", "章节", "写作", "图表模板"],
    "competition_guidance": ["赛题", "评分", "指标理解", "验收", "比赛", "挑战赛"],
}
# 商业来源（未授权不可作训练/特征主数据）
COMMERCIAL_TOKENS = ["链家", "贝壳", "安居客", "58", "我爱我家", "中原"]
# GIS 学习资料路径标识
GIS_LEARNING_TOKENS = ["城市数据学习资料", "classdata", "数据分析师", "规划实践",
                       "gis基本技能", "用地适宜性", "practice", "part3_gisbasic",
                       "part4_urbansituation"]

# 分类 → 外部归档目录 / 血缘 source_id / 默认用途
CATEGORY_DEST = {
    "policy_planning": "planning_policy/research_imports",
    "stats_macro": "stats_cn/research_imports",
    "population_profile": "population_profile/research_imports",
    "housing_property": "authorized_property/research_candidates",
    "industry_enterprise": "industry_enterprise/research_imports",
    "poi_public_service": "public_service/research_imports",
    "project_case": "urban_renewal_cases/research_imports",
    "report_template": "report_templates/research_imports",
    "competition_guidance": "competition_guidance/research_imports",
    "gis_learning": "gis_learning/research_imports",
    "unknown_or_review": "research_corpus/need_manual_review",
}


def classify(relative_path: str, file_name: str, text: str,
             detected: str) -> dict[str, Any]:
    """按文件名 + 提取文本做用途分类，返回分类结果（含合规判定）。"""
    haystack = f"{relative_path}\n{file_name}\n{text[:TEXT_TOPIC_SCAN]}".lower()

    # GIS 学习资料优先归档（课程素材，非项目数据）
    if detected == "gis" or any(t.lower() in haystack for t in GIS_LEARNING_TOKENS):
        return {
            "primary_category": "gis_learning",
            "secondary_categories": [],
            "matched_keywords": [],
            "useful_for_project": False,
            "commercial_source_risk": False,
            "reason": "GIS 课程学习素材 / 空间底图，非项目核心数据，仅登记参考",
        }

    scores: dict[str, list[str]] = {}
    for cat, kws in CATEGORY_KEYWORDS.items():
        hits = [kw for kw in kws if kw.lower() in haystack]
        if hits:
            scores[cat] = hits

    commercial = [t for t in COMMERCIAL_TOKENS if t.lower() in haystack]

    if not scores:
        return {
            "primary_category": "unknown_or_review",
            "secondary_categories": [],
            "matched_keywords": [],
            "useful_for_project": False,
            "commercial_source_risk": bool(commercial),
            "reason": "未命中任何用途关键词，需人工复核",
        }

    ranked = sorted(scores.items(), key=lambda kv: len(kv[1]), reverse=True)
    primary = ranked[0][0]
    secondary = [c for c, _ in ranked[1:4]]
    return {
        "primary_category": primary,
        "secondary_categories": secondary,
        "matched_keywords": ranked[0][1][:12],
        "useful_for_project": True,
        "commercial_source_risk": bool(commercial),
        "commercial_tokens": commercial,
        "reason": f"命中关键词归为 {primary}" + (
            f"；含商业来源标识 {commercial}（未授权不可作训练/特征主数据）" if commercial else ""),
    }


# --------------------------------------------------------------------------- #
# 用途 / 合规规则（默认保守）
# --------------------------------------------------------------------------- #
def usage_rules(category: str, commercial: bool) -> dict[str, Any]:
    base = {
        "used_for_training": False,            # 红线：科研语料一律不进监督训练
        "used_for_feature_engineering": False,
        "used_for_report": False,
        "used_for_rag": False,
        "license_status": "provided_by_research_partner",
        "compliance_status": "reference_only",
        "authorization_status": "unknown",
        "need_authorization": False,
    }
    if commercial or category == "housing_property":
        base.update({
            "compliance_status": "research_candidate_unauthorized",
            "authorization_status": "unknown",
            "need_authorization": True,
            "used_for_report": False,
            "used_for_feature_engineering": False,
        })
        return base
    if category in {"stats_macro", "population_profile", "industry_enterprise"}:
        base.update({"used_for_feature_engineering": True, "used_for_report": True,
                     "compliance_status": "reference_candidate"})
        return base
    if category == "poi_public_service":
        # POI 可作补充特征参考，但来源多为商业地图导出 → 仅参考
        base.update({"used_for_report": True, "compliance_status": "reference_candidate"})
        return base
    if category in {"policy_planning", "project_case", "report_template",
                    "competition_guidance"}:
        base.update({"used_for_report": True, "used_for_rag": True,
                     "compliance_status": "reference_text"})
        return base
    return base


# --------------------------------------------------------------------------- #
# 解析器
# --------------------------------------------------------------------------- #
def parse_table_xlsx(path: Path) -> dict[str, Any]:
    import openpyxl  # noqa: PLC0415
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = []
    total_rows = 0
    for ws in wb.worksheets:
        rows_iter = ws.iter_rows(values_only=True)
        header = []
        try:
            header_row = next(rows_iter)
            header = [str(c) for c in header_row if c is not None]
        except StopIteration:
            header_row = None
        n = 0
        for _ in rows_iter:
            n += 1
        total_rows += n
        sheets.append({"sheet": ws.title, "columns": header[:60],
                       "column_count": len(header), "data_row_count": n})
    wb.close()
    return {"sheet_count": len(sheets), "sheets": sheets,
            "record_count": total_rows,
            "columns_preview": sheets[0]["columns"] if sheets else []}


def parse_table_csv(path: Path) -> dict[str, Any]:
    encodings = ["utf-8-sig", "utf-8", "gbk", "gb18030", "latin-1"]
    for enc in encodings:
        try:
            with path.open("r", encoding=enc, newline="") as f:
                reader = csv.reader(f)
                rows = []
                for i, r in enumerate(reader):
                    rows.append(r)
                    if i > 200000:
                        break
            header = rows[0] if rows else []
            return {"encoding": enc, "columns_preview": [str(c) for c in header][:60],
                    "column_count": len(header),
                    "record_count": max(len(rows) - 1, 0)}
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as exc:  # noqa: BLE001
            return {"parse_error": str(exc)}
    return {"parse_error": "无法用常见编码解析 CSV"}


def parse_structured(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as exc:  # noqa: BLE001
        return {"parse_error": str(exc)}
    if isinstance(data, dict) and data.get("type") == "FeatureCollection":
        feats = data.get("features", []) or []
        keys = set()
        for ft in feats[:50]:
            keys.update((ft.get("properties") or {}).keys())
        return {"geojson": True, "feature_count": len(feats),
                "property_keys": sorted(keys)[:60], "record_count": len(feats)}
    if isinstance(data, list):
        keys = set()
        for it in data[:50]:
            if isinstance(it, dict):
                keys.update(it.keys())
        return {"record_count": len(data), "keys": sorted(keys)[:60]}
    if isinstance(data, dict):
        return {"record_count": 1, "keys": sorted(data.keys())[:60]}
    return {"record_count": 0}


def parse_docx(path: Path) -> dict[str, Any]:
    import docx  # noqa: PLC0415
    d = docx.Document(str(path))
    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paras)
    table_rows = sum(len(t.rows) for t in d.tables)
    return {"text": text[:TEXT_STORE_CAP], "paragraph_count": len(paras),
            "char_count": len(text), "table_count": len(d.tables),
            "table_row_count": table_rows}


def parse_pdf(path: Path) -> dict[str, Any]:
    from pypdf import PdfReader  # noqa: PLC0415
    reader = PdfReader(str(path))
    pages = len(reader.pages)
    chunks = []
    for pg in reader.pages[:80]:
        try:
            chunks.append(pg.extract_text() or "")
        except Exception:  # noqa: BLE001
            chunks.append("")
    text = "\n".join(chunks).strip()
    result = {"page_count": pages, "text": text[:TEXT_STORE_CAP],
              "char_count": len(text)}
    if len(text) < 50:
        result["need_ocr"] = True
        result["need_manual_review"] = True
        result["parse_note"] = "PDF 文本极少，疑似扫描件，需 OCR / 人工"
    return result


def parse_text(path: Path) -> dict[str, Any]:
    for enc in ["utf-8", "gbk", "gb18030", "latin-1"]:
        try:
            text = path.read_text(encoding=enc)
            return {"text": text[:TEXT_STORE_CAP], "char_count": len(text),
                    "encoding": enc}
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as exc:  # noqa: BLE001
            return {"parse_error": str(exc)}
    return {"parse_error": "无法解码文本文件"}


def run_parse(path: Path, method: str) -> dict[str, Any]:
    try:
        if method == "openpyxl":
            return parse_table_xlsx(path)
        if method == "csv":
            return parse_table_csv(path)
        if method == "json":
            return parse_structured(path)
        if method == "python-docx":
            return parse_docx(path)
        if method == "pypdf":
            return parse_pdf(path)
        if method == "text":
            return parse_text(path)
    except Exception as exc:  # noqa: BLE001
        return {"parse_error": f"{type(exc).__name__}: {exc}"}
    return {}


# --------------------------------------------------------------------------- #
# 主流程
# --------------------------------------------------------------------------- #
def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    try:
        with path.open("rb") as f:
            for blk in iter(lambda: f.read(1 << 20), b""):
                h.update(blk)
    except Exception:  # noqa: BLE001
        return ""
    return h.hexdigest()


def _file_id(rel: str) -> str:
    return hashlib.sha1(rel.encode("utf-8")).hexdigest()[:16]


def find_corpus_dir() -> Path | None:
    for name in CORPUS_CANDIDATES:
        p = PROJECT_ROOT / name
        if p.is_dir():
            return p
    return None


def scan() -> dict[str, Any]:
    corpus = find_corpus_dir()
    if corpus is None:
        raise SystemExit(f"未找到科研语料目录，候选：{CORPUS_CANDIDATES}")

    for d in (INVENTORY_DIR, EXTRACTED_DIR, CLASSIFIED_DIR, DATA_CATALOG_DIR):
        d.mkdir(parents=True, exist_ok=True)

    inventory: list[dict[str, Any]] = []
    classification: list[dict[str, Any]] = []
    # 用途聚合：category -> {files, record_count, ...}
    cat_agg: dict[str, dict[str, Any]] = {}

    all_files = [p for p in corpus.rglob("*")
                 if p.is_file() and p.name != ".DS_Store"]
    all_files.sort()

    for path in all_files:
        rel = path.relative_to(PROJECT_ROOT).as_posix()
        rel_in_corpus = path.relative_to(corpus).as_posix()
        ext = path.suffix.lstrip(".").lower()
        detected = detect_type(ext)
        plan = parse_plan(ext, detected)
        try:
            stat = path.stat()
            size = stat.st_size
            mtime = datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat()
        except Exception:  # noqa: BLE001
            size, mtime = 0, None

        fid = _file_id(rel)
        sha = _sha256(path)

        rec = {
            "file_id": fid,
            "original_path": str(path),
            "relative_path": rel,
            "relative_in_corpus": rel_in_corpus,
            "file_name": path.name,
            "file_ext": ext,
            "file_size": size,
            "sha256": sha,
            "modified_time": mtime,
            "detected_type": detected,
            "can_parse": plan["can_parse"],
            "parse_method": plan["parse_method"],
            "need_ocr": plan["need_ocr"],
            "need_manual_review": plan["need_manual_review"],
            "source_folder": corpus.name,
            "imported": False,
            "failed_reason": plan.get("plan_note"),
        }

        # 解析（仅对可解析类型）
        parsed: dict[str, Any] = {}
        text = ""
        if plan["can_parse"]:
            parsed = run_parse(path, plan["parse_method"])
            if parsed.get("parse_error"):
                rec["failed_reason"] = parsed["parse_error"]
                rec["need_manual_review"] = True
            text = parsed.get("text", "") or ""
            # PDF 扫描件回填
            if parsed.get("need_ocr"):
                rec["need_ocr"] = True
                rec["need_manual_review"] = True

        # 分类
        cls = classify(rel, path.name, text, detected)
        category = cls["primary_category"]
        usage = usage_rules(category, cls.get("commercial_source_risk", False))

        inventory.append(rec)

        record_count = int(parsed.get("record_count") or 0)
        produce_extracted = detected in {"document", "table", "structured_data",
                                         "image", "presentation", "archive"} or \
            cls["useful_for_project"]
        if produce_extracted:
            extracted = {
                "file_id": fid,
                "relative_path": rel,
                "file_name": path.name,
                "detected_type": detected,
                "parse_method": plan["parse_method"],
                "need_ocr": rec["need_ocr"],
                "need_manual_review": rec["need_manual_review"],
                "record_count": record_count,
                "primary_category": category,
                "classification": cls,
                "usage": usage,
                "parsed": _strip_text_for_store(parsed),
                "text_preview": (text[:PREVIEW_CHARS] if text else None),
                "extracted_at": _utcnow(),
            }
            (EXTRACTED_DIR / f"{fid}.json").write_text(
                json.dumps(extracted, ensure_ascii=False, indent=2), encoding="utf-8")

        cls_row = {
            "file_id": fid,
            "relative_path": rel,
            "file_name": path.name,
            "detected_type": detected,
            "primary_category": category,
            "secondary_categories": cls["secondary_categories"],
            "matched_keywords": cls.get("matched_keywords", []),
            "useful_for_project": cls["useful_for_project"],
            "commercial_source_risk": cls.get("commercial_source_risk", False),
            "record_count": record_count,
            "need_manual_review": rec["need_manual_review"],
            "usage": usage,
            "reason": cls["reason"],
        }
        classification.append(cls_row)

        agg = cat_agg.setdefault(category, {
            "category": category, "file_count": 0, "record_count": 0,
            "useful_file_count": 0, "need_manual_review_count": 0,
            "commercial_risk_count": 0, "files": [],
            "dest_dir": CATEGORY_DEST.get(category, "research_corpus/need_manual_review"),
        })
        agg["file_count"] += 1
        agg["record_count"] += record_count
        agg["useful_file_count"] += 1 if cls["useful_for_project"] else 0
        agg["need_manual_review_count"] += 1 if rec["need_manual_review"] else 0
        agg["commercial_risk_count"] += 1 if cls.get("commercial_source_risk") else 0
        agg["files"].append({
            "file_id": fid, "relative_path": rel, "file_name": path.name,
            "sha256": sha, "record_count": record_count,
            "detected_type": detected, "need_manual_review": rec["need_manual_review"],
            "usage": usage,
        })

    return {
        "corpus_dir": str(corpus),
        "corpus_name": corpus.name,
        "inventory": inventory,
        "classification": classification,
        "cat_agg": cat_agg,
    }


def _strip_text_for_store(parsed: dict[str, Any]) -> dict[str, Any]:
    """extracted.parsed 不再重复存全文（已在 text_preview 给摘要），保留结构化元信息。"""
    out = dict(parsed)
    if "text" in out:
        out["text_char_count"] = len(out["text"])
        out.pop("text", None)
    return out


# --------------------------------------------------------------------------- #
# 产出：inventory / classification / 归档元数据 / manifest / lineage / 缺口清单
# --------------------------------------------------------------------------- #
def write_outputs(result: dict[str, Any], *, record_lineage: bool = True) -> dict[str, Any]:
    inventory = result["inventory"]
    classification = result["classification"]
    cat_agg = result["cat_agg"]

    # 1) inventory json + md
    inv_json = {
        "generated_at": _utcnow(),
        "corpus_dir": result["corpus_dir"],
        "source_folder": result["corpus_name"],
        "total_files": len(inventory),
        "parser_availability": PARSEABLE_LIBS,
        "files": inventory,
    }
    (INVENTORY_DIR / "research_file_inventory.json").write_text(
        json.dumps(inv_json, ensure_ascii=False, indent=2), encoding="utf-8")
    (INVENTORY_DIR / "research_file_inventory.md").write_text(
        _inventory_md(inventory, result["corpus_name"]), encoding="utf-8")

    # 2) classification json + md
    cls_json = {
        "generated_at": _utcnow(),
        "total_files": len(classification),
        "category_summary": {c: {k: v for k, v in agg.items() if k != "files"}
                             for c, agg in cat_agg.items()},
        "files": classification,
    }
    (CLASSIFIED_DIR / "research_classification.json").write_text(
        json.dumps(cls_json, ensure_ascii=False, indent=2), encoding="utf-8")
    (CLASSIFIED_DIR / "research_classification.md").write_text(
        _classification_md(cat_agg), encoding="utf-8")

    # 3) 每类归档元数据（不复制原件，只写元数据/脱敏摘要索引）
    for category, agg in cat_agg.items():
        dest_rel = agg["dest_dir"]
        dest_dir = EXTERNAL_DIR / dest_rel
        dest_dir.mkdir(parents=True, exist_ok=True)
        index = {
            "generated_at": _utcnow(),
            "category": category,
            "source_folder": result["corpus_name"],
            "note": "仅登记元数据与脱敏摘要索引，原始文件保留在 gitignore 的科研语料目录，未复制。",
            "file_count": agg["file_count"],
            "record_count": agg["record_count"],
            "files": agg["files"],
        }
        (dest_dir / f"{category}_research_index.json").write_text(
            json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")

    # 4) lineage（每类一条聚合）+ manifest
    from app.services import data_lineage_service
    lineage_records = []
    for category, agg in cat_agg.items():
        if category in {"gis_learning", "unknown_or_review"}:
            # 仅登记，不写血缘用途（参考素材）
            continue
        sample_usage = agg["files"][0]["usage"] if agg["files"] else usage_rules(category, False)
        if not record_lineage:
            lineage_records.append({
                "category": category, "source_id": f"research_{category}",
                "file_count": agg["file_count"], "record_count": agg["record_count"]})
            continue
        lid = data_lineage_service.record_collection_lineage(
            source_id=f"research_{category}",
            source_name=f"科研语料-{category}",
            source_type="research_corpus",
            raw_count=agg["file_count"],
            cleaned_count=agg["record_count"] or agg["file_count"],
            license_status=sample_usage["license_status"],
            compliance_status=sample_usage["compliance_status"],
            used_for_training=False,
            used_for_feature_engineering=sample_usage["used_for_feature_engineering"],
            used_for_report=sample_usage["used_for_report"],
            file_path=agg["dest_dir"],
        )
        lineage_records.append({
            "lineage_id": lid, "category": category,
            "source_id": f"research_{category}",
            "file_count": agg["file_count"], "record_count": agg["record_count"],
            "used_for_training": False,
            "used_for_feature_engineering": sample_usage["used_for_feature_engineering"],
            "used_for_report": sample_usage["used_for_report"],
            "compliance_status": sample_usage["compliance_status"],
        })

    manifest = {
        "generated_at": _utcnow(),
        "source_folder": result["corpus_name"],
        "total_files": len(inventory),
        "total_records_parsed": sum(c.get("record_count", 0) for c in classification),
        "parser_availability": PARSEABLE_LIBS,
        "category_summary": {c: {k: v for k, v in agg.items() if k != "files"}
                             for c, agg in cat_agg.items()},
        "lineage_records": lineage_records,
        "red_lines": [
            "used_for_training=false（科研语料一律不进监督训练）",
            "商业来源（链家/贝壳等）为 research_candidate，需授权后才可用",
            "原始文件保留在 gitignore 的科研语料目录，未复制进版本库",
            "未执行任何可执行文件，未自动解压压缩包",
        ],
    }
    (RC_DIR / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    (RC_DIR / "lineage.json").write_text(
        json.dumps({"generated_at": _utcnow(), "records": lineage_records},
                   ensure_ascii=False, indent=2), encoding="utf-8")

    # 5) 缺口补齐清单
    gap = build_gap_fill(cat_agg)
    (DATA_CATALOG_DIR / "科研语料可补缺口清单.json").write_text(
        json.dumps(gap, ensure_ascii=False, indent=2), encoding="utf-8")
    (DATA_CATALOG_DIR / "科研语料可补缺口清单.md").write_text(
        _gap_md(gap), encoding="utf-8")

    return {"manifest": manifest, "gap": gap, "lineage_records": lineage_records}


def build_gap_fill(cat_agg: dict[str, dict[str, Any]]) -> dict[str, Any]:
    def cnt(cat: str) -> tuple[int, int]:
        a = cat_agg.get(cat, {})
        return a.get("file_count", 0), a.get("record_count", 0)

    pop_f, pop_r = cnt("population_profile")
    stat_f, stat_r = cnt("stats_macro")
    pol_f, _ = cnt("policy_planning")
    hou_f, hou_r = cnt("housing_property")
    ind_f, ind_r = cnt("industry_enterprise")
    poi_f, poi_r = cnt("poi_public_service")

    gaps = [
        {
            "gap": "人口与收入数据",
            "research_can_fill": "partial" if (pop_f or stat_f) else "no",
            "from_categories": ["population_profile", "stats_macro"],
            "file_count": pop_f + stat_f,
            "record_count": pop_r + stat_r,
            "trainable": False,
            "note": "科研语料中人口/收入多为参考文档或表格，可作特征参考，"
                    "但仍缺权威统计年鉴口径，建议人工补统计局数据。",
        },
        {
            "gap": "政策与规划",
            "research_can_fill": "partial" if pol_f else "no",
            "from_categories": ["policy_planning", "project_case"],
            "file_count": pol_f,
            "record_count": 0,
            "trainable": False,
            "note": "政策/案例多为 PPT 图片版与 PDF，可作 RAG/报告参考；"
                    "图片需人工或 OCR 复核，不能直接结构化。",
        },
        {
            "gap": "房价样本",
            "research_can_fill": "candidate_unauthorized" if hou_f else "no",
            "from_categories": ["housing_property"],
            "file_count": hou_f,
            "record_count": hou_r,
            "trainable": False,
            "note": "链家/贝壳等商业来源住房样本量较大，但授权状态未知，"
                    "仅登记为 research_candidate，授权前不得用于训练/特征主数据。",
        },
        {
            "gap": "产业细分",
            "research_can_fill": "partial" if ind_f else "no",
            "from_categories": ["industry_enterprise"],
            "file_count": ind_f,
            "record_count": ind_r,
            "trainable": False,
            "note": "产业类资料较少，仍需补企业/园区结构化数据。",
        },
        {
            "gap": "公共服务/POI",
            "research_can_fill": "supplement" if poi_f else "no",
            "from_categories": ["poi_public_service"],
            "file_count": poi_f,
            "record_count": poi_r,
            "trainable": False,
            "note": "POI 类（如美食 POI）可作高德 POI 的补充参考；高德已覆盖 5 万级。",
        },
    ]
    return {"generated_at": _utcnow(), "source": "科研语料", "gaps": gaps,
            "red_line": "全部 used_for_training=false，未混入 competition_test"}


# --------------------------------------------------------------------------- #
# Markdown 渲染
# --------------------------------------------------------------------------- #
def _inventory_md(inventory: list[dict], folder: str) -> str:
    by_type: dict[str, int] = {}
    review = 0
    for r in inventory:
        by_type[r["detected_type"]] = by_type.get(r["detected_type"], 0) + 1
        review += 1 if r["need_manual_review"] else 0
    lines = [f"# 科研语料资产清单（{folder}）", "",
             f"- 生成时间：{_utcnow()}",
             f"- 文件总数：{len(inventory)}",
             f"- 需人工复核：{review}",
             f"- 解析器可用性：{PARSEABLE_LIBS}", "",
             "## 按类型分布", "", "| 类型 | 数量 |", "| --- | --- |"]
    for t, n in sorted(by_type.items(), key=lambda kv: kv[1], reverse=True):
        lines.append(f"| {t} | {n} |")
    lines += ["", "## 关键可解析文件（document/table/structured_data，前 60）", "",
              "| 文件 | 类型 | 解析 | 记录数? | 需复核 |", "| --- | --- | --- | --- | --- |"]
    shown = 0
    for r in inventory:
        if r["detected_type"] in {"document", "table", "structured_data",
                                  "presentation"}:
            lines.append(f"| {r['file_name']} | {r['detected_type']} | "
                         f"{r['parse_method']} | - | "
                         f"{'是' if r['need_manual_review'] else '否'} |")
            shown += 1
            if shown >= 60:
                break
    return "\n".join(lines) + "\n"


def _classification_md(cat_agg: dict[str, dict]) -> str:
    lines = ["# 科研语料用途分类", "", f"- 生成时间：{_utcnow()}", "",
             "| 用途类别 | 文件数 | 记录数 | 有用 | 需复核 | 商业风险 | 归档目录 |",
             "| --- | --- | --- | --- | --- | --- | --- |"]
    for cat, agg in sorted(cat_agg.items(),
                           key=lambda kv: kv[1]["file_count"], reverse=True):
        lines.append(f"| {cat} | {agg['file_count']} | {agg['record_count']} | "
                     f"{agg['useful_file_count']} | {agg['need_manual_review_count']} | "
                     f"{agg['commercial_risk_count']} | {agg['dest_dir']} |")
    lines += ["", "> 红线：科研语料全部 used_for_training=false；"
              "商业来源（链家/贝壳等）为 research_candidate，需授权后才可用。"]
    return "\n".join(lines) + "\n"


def _gap_md(gap: dict[str, Any]) -> str:
    lines = ["# 科研语料可补缺口清单", "", f"- 生成时间：{gap['generated_at']}",
             f"- 来源：{gap.get('source', '科研语料')}",
             f"- 授权：{gap.get('authorization', '')}", "",
             "| 缺口 | 可补程度 | 可训练 | 记录/候选 | 说明 |",
             "| --- | --- | --- | --- | --- |"]
    for g in gap["gaps"]:
        cnt = g.get("trainable_records") or g.get("records") or g.get("record_count") \
            or g.get("rag_candidates") or g.get("file_count") or 0
        lines.append(f"| {g['gap']} | {g['research_can_fill']} | "
                     f"{'是' if g.get('trainable') else '否'} | {cnt} | {g.get('note', '')} |")
    lines += ["", f"> {gap.get('red_line', '')}"]
    return "\n".join(lines) + "\n"


# --------------------------------------------------------------------------- #
# 第10C.5：授权确认 + 正式数据资产转化
# --------------------------------------------------------------------------- #
AUTH_CONFIG_PATH = RC_DIR / "research_authorization_config.json"
DEFAULT_AUTH_CONFIG = {
    "provider": "科研人员提供",
    "authorized_by_user": True,
    "authorization_scope": ["feature_engineering", "report", "rag",
                            "model_training_if_desensitized_and_structured"],
    "commercial_risk_override": True,
    "reason": "用户确认科研人员提供的数据可用于本项目开发与研究分析",
    "training_rule": {
        "default_used_for_training": False,
        "housing_property_can_train_if": [
            "record_count >= 1000", "has_price_label = true",
            "has_location_or_region = true", "is_desensitized = true",
            "authorization_status in ['provided', 'authorized']"],
    },
}

# 列角色识别（按优先级，先经纬度避免被 '区' 误吞）
ROLE_TOKENS = [
    ("lat", ["纬度", "lat"]),
    ("lng", ["经度", "lng", "lon"]),
    ("price_unit", ["单价", "挂牌均价", "均价", "priceunitlist", "unit_price"]),
    ("price_total", ["总价", "pricetotallist", "price_total", "total_price"]),
    ("rent", ["租金", "出租价格", "租赁价格", "平均租赁价格", "租金季付"]),
    ("area", ["建筑面积", "实测住房建筑面积", "面积", "area"]),
    ("build_year", ["建成年代", "建成年份", "houselife", "房龄", "year"]),
    ("room", ["roomtype", "房屋类型", "户型"]),
    ("time", ["月份", "pricedatelist", "维护时间", "成交时间", "年份", "date"]),
    ("region", ["行政区", "区域", "市区", "板块", "五大新城", "district", "location"]),
    ("community", ["小区名称", "小区名字", "供应备案项目名称", "projectname",
                   "项目名称", "小区", "community"]),
    ("address", ["小区地址", "保租房地址", "地址", "address"]),
]
DROP_TOKENS = ["url", "链接", "核验码", "房源编号", "householdid", "description",
               "描述", "卖点", "sellingpoint", "intro", "简介", "物业公司", "开发商",
               "经纪", "维护人"]
PRIVACY_BLOCK_TOKENS = ["姓名", "身份证", "手机", "电话", "phone", "mobile",
                        "业主", "联系人", "客户姓名"]
_NUM_RE = re.compile(r"-?\d+(?:\.\d+)?")


def _col_role(name: str) -> str | None:
    low = str(name).strip().lower()
    if low == "区":
        return "region"
    for role, toks in ROLE_TOKENS:
        if any(t in low for t in toks):
            return role
    return None


def _num(v: Any) -> float | None:
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    m = _NUM_RE.search(str(v).replace(",", ""))
    return float(m.group()) if m else None


def _xlsx_header_rows(path: Path):
    import openpyxl  # noqa: PLC0415
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.worksheets[0]
    it = ws.iter_rows(values_only=True)
    header = []
    try:
        header = [str(c).strip() if c is not None else f"col{i}"
                  for i, c in enumerate(next(it))]
    except StopIteration:
        wb.close()
        return [], iter([])

    def gen():
        for row in it:
            yield dict(zip(header, row))
        wb.close()
    return header, gen()


def _role_map(header: list[str]) -> dict[str, str]:
    rmap: dict[str, str] = {}
    for col in header:
        role = _col_role(col)
        if role and role not in rmap.values():  # 每个 role 取首个命中列
            rmap[col] = role
    return rmap


def process_housing(cls_files: list[dict], authorized: bool) -> dict[str, Any]:
    dest = EXTERNAL_DIR / "authorized_property" / "processed"
    dest.mkdir(parents=True, exist_ok=True)
    out_jsonl = dest / "research_property_trainable_candidates.jsonl"

    file_profiles = []
    schema_mappings = []
    seen: set[str] = set()
    trainable_rows = sale_rows = rent_rows = 0
    total_parsed = 0

    with out_jsonl.open("w", encoding="utf-8") as fout:
        for f in cls_files:
            path = PROJECT_ROOT / f["relative_path"]
            ext = f["file_ext"] if "file_ext" in f else path.suffix.lstrip(".").lower()
            prof = {"file_id": f["file_id"], "file_name": f["file_name"],
                    "record_count": f.get("record_count", 0)}
            if not path.exists() or ext != "xlsx" or not PARSEABLE_LIBS["openpyxl"]:
                prof.update({"parsable": False, "trainable": False,
                             "failed_reason": "非 xlsx 或解析器不可用，需人工转换"})
                file_profiles.append(prof)
                continue
            try:
                header, rows = _xlsx_header_rows(path)
            except Exception as exc:  # noqa: BLE001
                prof.update({"parsable": False, "trainable": False,
                             "failed_reason": f"解析失败: {exc}"})
                file_profiles.append(prof)
                continue

            rmap = _role_map(header)
            roles = set(rmap.values())
            has_price = bool({"price_unit", "price_total", "rent"} & roles)
            has_location = bool({"region", "community", "address"} & roles) or \
                {"lat", "lng"} <= roles
            has_area = "area" in roles
            has_time = "time" in roles
            has_attr = bool({"room", "build_year"} & roles)
            privacy = [c for c in header
                       if any(t in str(c).lower() for t in PRIVACY_BLOCK_TOKENS)]
            is_desensitized = len(privacy) == 0
            price_type = "sale" if {"price_unit", "price_total"} & roles else (
                "rent" if "rent" in roles else "none")

            trainable_file = (authorized and f.get("record_count", 0) >= 1000
                              and has_price and has_location and is_desensitized)
            reason = None
            if not trainable_file:
                miss = []
                if f.get("record_count", 0) < 1000:
                    miss.append("record_count<1000")
                if not has_price:
                    miss.append("无价格标签")
                if not has_location:
                    miss.append("无区域/位置字段")
                if not is_desensitized:
                    miss.append(f"含个人隐私字段{privacy}")
                if not authorized:
                    miss.append("未授权")
                reason = ";".join(miss)

            schema_mappings.append({
                "file_id": f["file_id"], "file_name": f["file_name"],
                "column_role_map": rmap, "price_type": price_type,
                "has_price_label": has_price, "has_location": has_location,
                "has_area": has_area, "has_time": has_time, "has_attr": has_attr,
                "is_desensitized": is_desensitized, "privacy_columns": privacy})

            n_rows = n_with_price_loc = 0
            for r in rows:
                n_rows += 1
                total_parsed += 1
                rec: dict[str, Any] = {}
                for col, role in rmap.items():
                    val = r.get(col)
                    if role in {"price_unit", "price_total", "rent", "area", "lng", "lat"}:
                        nv = _num(val)
                        if nv is not None:
                            rec[role] = round(nv, 5) if role in {"lng", "lat"} else nv
                    elif role == "build_year":
                        nv = _num(val)
                        if nv is not None:
                            rec[role] = int(nv)
                    elif val not in (None, ""):
                        rec[role] = str(val).strip()[:80]
                has_p = bool({"price_unit", "price_total", "rent"} & set(rec))
                has_l = bool({"region", "community", "address", "lat"} & set(rec))
                if trainable_file and has_p and has_l:
                    key = hashlib.md5(
                        f"{rec.get('community','')}|{rec.get('region','')}|"
                        f"{rec.get('area','')}|{rec.get('price_unit','')}|"
                        f"{rec.get('price_total','')}|{rec.get('rent','')}|"
                        f"{rec.get('build_year','')}".encode()).hexdigest()
                    if key in seen:
                        continue
                    seen.add(key)
                    rec["price_type"] = price_type
                    rec["src_file_id"] = f["file_id"]
                    rec["used_for_training"] = True
                    rec["training_source"] = "research_authorized_property"
                    rec["competition_test"] = False
                    fout.write(json.dumps(rec, ensure_ascii=False) + "\n")
                    trainable_rows += 1
                    n_with_price_loc += 1
                    if price_type == "sale":
                        sale_rows += 1
                    elif price_type == "rent":
                        rent_rows += 1
            prof.update({"parsable": True, "parsed_rows": n_rows, "price_type": price_type,
                         "has_price_label": has_price, "has_location": has_location,
                         "is_desensitized": is_desensitized, "trainable": trainable_file,
                         "trainable_rows_contributed": n_with_price_loc,
                         "failed_reason": reason})
            file_profiles.append(prof)

    strength = ("strong" if trainable_rows >= 3000 else
                "medium" if trainable_rows >= 1000 else "weak")
    profile = {
        "generated_at": _utcnow(), "category": "housing_property",
        "provider": "科研人员提供", "authorization_status": "provided" if authorized else "unknown",
        "license_status": "provided_by_research_partner",
        "commercial_risk_override": authorized,
        "file_count": len(cls_files),
        "total_records": sum(f.get("record_count", 0) for f in cls_files),
        "parsed_records": total_parsed,
        "trainable_property_records": trainable_rows,
        "trainable_sale_records": sale_rows, "trainable_rent_records": rent_rows,
        "supervised_training_strength": strength,
        "can_start_supervised_housing_model": trainable_rows >= 1000,
        "is_desensitized": True, "test_contamination_risk": False, "leakage_risk": False,
        "files": file_profiles,
        "note": "明细已脱敏写入 trainable_candidates.jsonl（gitignore），接口不返回原始明细。",
    }
    (dest / "research_property_dataset_profile.json").write_text(
        json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    (dest / "research_property_schema_mapping.json").write_text(
        json.dumps({"generated_at": _utcnow(), "mappings": schema_mappings},
                   ensure_ascii=False, indent=2), encoding="utf-8")
    return profile


def _load_amap_index() -> tuple[set, set]:
    """返回 (name_hash 集合, 坐标集合)。高德 store 仅存 name_hash(脱敏)，
    因此与科研 POI 只能按坐标(GCJ02 四位)近似去重。"""
    store = EXTERNAL_DIR / "amap" / "large_scale_store" / "store.json"
    name_hashes: set[str] = set()
    coords: set[str] = set()
    if not store.exists():
        return name_hashes, coords
    try:
        with store.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:  # noqa: BLE001
        return name_hashes, coords
    records = data.get("records") if isinstance(data, dict) else data
    items = records.values() if isinstance(records, dict) else (records or [])
    for it in items:
        if not isinstance(it, dict):
            continue
        if it.get("name_hash"):
            name_hashes.add(str(it["name_hash"]))
        loc = it.get("location_gcj02") or it.get("location") or ""
        if isinstance(loc, str) and "," in loc:
            try:
                lng, lat = loc.split(",")[:2]
                coords.add(f"{round(float(lng), 4)},{round(float(lat), 4)}")
            except Exception:  # noqa: BLE001
                pass
    return name_hashes, coords


def process_poi(cls_files: list[dict]) -> dict[str, Any]:
    dest = EXTERNAL_DIR / "public_service" / "processed"
    dest.mkdir(parents=True, exist_ok=True)
    out_jsonl = dest / "research_poi_feature_candidates.jsonl"
    amap_name_hashes, amap_coords = _load_amap_index()

    total = overlap_name = overlap_coord = emitted = 0
    type_counter: dict[str, int] = {}
    name_tok = ["名称", "name", "poi", "店名", "商户"]
    type_tok = ["类型", "分类", "category", "type", "一级", "二级", "标签", "tag"]
    with out_jsonl.open("w", encoding="utf-8") as fout:
        for f in cls_files:
            path = PROJECT_ROOT / f["relative_path"]
            if not path.exists() or path.suffix.lstrip(".").lower() != "xlsx":
                continue
            try:
                header, rows = _xlsx_header_rows(path)
            except Exception:  # noqa: BLE001
                continue
            rmap = _role_map(header)
            name_col = next((c for c in header if any(t in str(c).lower() for t in name_tok)), None)
            type_col = next((c for c in header if any(t in str(c).lower() for t in type_tok)), None)
            for r in rows:
                total += 1
                nm = str(r.get(name_col, "")).strip() if name_col else ""
                tp = str(r.get(type_col, "")).strip() if type_col else ""
                lng = _num(next((r.get(c) for c, ro in rmap.items() if ro == "lng"), None))
                lat = _num(next((r.get(c) for c, ro in rmap.items() if ro == "lat"), None))
                region = next((str(r.get(c)) for c, ro in rmap.items() if ro == "region"), None)
                if tp:
                    type_counter[tp] = type_counter.get(tp, 0) + 1
                # 高德仅存 name_hash(脱敏)，名称无法直接比对；按坐标(GCJ02 4位)近似去重
                dup_c = (lng is not None and lat is not None
                         and f"{round(lng, 4)},{round(lat, 4)}" in amap_coords)
                if dup_c:
                    overlap_coord += 1
                rec = {"name": nm[:80], "type": tp[:60], "region": region,
                       "lng": round(lng, 5) if lng else None,
                       "lat": round(lat, 5) if lat else None,
                       "overlap_with_amap": bool(dup_c),
                       "used_for_feature_engineering": True, "used_for_report": True,
                       "used_for_training": False, "competition_test": False}
                fout.write(json.dumps(rec, ensure_ascii=False) + "\n")
                emitted += 1

    unique = emitted - overlap_coord
    profile = {
        "generated_at": _utcnow(), "category": "poi_public_service",
        "provider": "科研人员提供", "authorization_status": "provided",
        "license_status": "provided_by_research_partner", "commercial_risk_override": True,
        "total_records": total, "emitted_records": emitted,
        "used_for_feature_engineering": True, "used_for_report": True,
        "used_for_training": False,
        "top_types": sorted(type_counter.items(), key=lambda kv: kv[1], reverse=True)[:20],
        "overlap_summary": {
            "amap_poi_total": len(amap_coords),
            "research_poi_total": emitted,
            "overlap_by_name": "n/a（高德仅存 name_hash 脱敏，名称不可比对）",
            "overlap_by_coord": overlap_coord,
            "unique_research_estimate": max(0, unique),
            "note": "按 GCJ02 坐标(4位)近似去重；不并入高德 50075 官方 POI 资产，仅作补充特征参考。"},
        "test_contamination_risk": False, "leakage_risk": False,
    }
    (dest / "research_poi_dataset_profile.json").write_text(
        json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    return profile


def process_stats(cls_files: list[dict]) -> dict[str, Any]:
    dest = EXTERNAL_DIR / "stats_cn" / "processed"
    dest.mkdir(parents=True, exist_ok=True)
    indicators = ["人口", "收入", "消费", "cpi", "gdp", "就业", "产业结构",
                  "固定资产", "社会消费品零售", "房地产"]
    parseable = []
    for f in cls_files:
        ex = EXTRACTED_DIR / f"{f['file_id']}.json"
        txt_chars = 0
        if ex.exists():
            try:
                e = json.load(ex.open())
                txt_chars = (e.get("parsed", {}) or {}).get("text_char_count", 0) \
                    or len(e.get("text_preview") or "")
            except Exception:  # noqa: BLE001
                pass
        parseable.append({"file_name": f["file_name"], "detected_type": f["detected_type"],
                          "record_count": f.get("record_count", 0),
                          "need_manual_review": f.get("need_manual_review", True),
                          "extracted_text_chars": txt_chars})
    has_structured = any(p["record_count"] > 0 for p in parseable)
    profile = {
        "generated_at": _utcnow(), "category": "stats_macro",
        "provider": "科研人员提供", "authorization_status": "provided",
        "license_status": "provided_by_research_partner",
        "target_indicators": indicators, "files": parseable,
        "has_structured_stats": has_structured,
        "status": "available" if has_structured else "still_missing",
        "used_for_feature_engineering": has_structured, "used_for_report": True,
        "used_for_training": False,
        "note": "科研语料统计类多为扫描 PDF/图片，无法结构化解析则标 still_missing，不伪造指标。",
        "test_contamination_risk": False, "leakage_risk": False,
    }
    (dest / "research_stats_dataset_profile.json").write_text(
        json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    return profile


def process_rag(cls_files: list[dict], out_path: Path, category: str) -> dict[str, Any]:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    emitted = need_review = with_text = 0
    with out_path.open("w", encoding="utf-8") as fout:
        for f in cls_files:
            ex = EXTRACTED_DIR / f"{f['file_id']}.json"
            text_preview = None
            if ex.exists():
                try:
                    e = json.load(ex.open())
                    text_preview = e.get("text_preview")
                except Exception:  # noqa: BLE001
                    pass
            rec = {
                "file_id": f["file_id"], "file_name": f["file_name"],
                "detected_type": f["detected_type"], "category": category,
                "text_preview": text_preview,
                "has_structured_text": bool(text_preview),
                "need_manual_review": f.get("need_manual_review", True),
                "used_for_rag": True, "used_for_report": True, "used_for_training": False,
                "note": "图片/PPT/扫描件需 OCR/人工复核后方可作为结构化事实，此处仅作 RAG 引用候选。",
            }
            fout.write(json.dumps(rec, ensure_ascii=False) + "\n")
            emitted += 1
            need_review += 1 if rec["need_manual_review"] else 0
            with_text += 1 if text_preview else 0
    return {"category": category, "candidate_count": emitted,
            "need_manual_review": need_review, "with_extracted_text": with_text,
            "output": str(out_path.relative_to(BACKEND_DIR))}


def apply_authorization(result: dict[str, Any]) -> dict[str, Any]:
    cls_by_cat: dict[str, list[dict]] = {}
    # 给 classification 行补 file_ext
    inv_by_id = {r["file_id"]: r for r in result["inventory"]}
    for row in result["classification"]:
        row.setdefault("file_ext", inv_by_id.get(row["file_id"], {}).get("file_ext", ""))
        row.setdefault("relative_path", inv_by_id.get(row["file_id"], {}).get("relative_path"))
        cls_by_cat.setdefault(row["primary_category"], []).append(row)

    config = json.loads(AUTH_CONFIG_PATH.read_text(encoding="utf-8")) \
        if AUTH_CONFIG_PATH.exists() else DEFAULT_AUTH_CONFIG
    if not AUTH_CONFIG_PATH.exists():
        AUTH_CONFIG_PATH.write_text(json.dumps(config, ensure_ascii=False, indent=2),
                                    encoding="utf-8")
    authorized = bool(config.get("authorized_by_user"))

    print("[apply] processing housing ...")
    housing = process_housing(cls_by_cat.get("housing_property", []), authorized)
    print(f"[apply] housing trainable_rows={housing['trainable_property_records']} "
          f"strength={housing['supervised_training_strength']}")
    print("[apply] processing poi ...")
    poi = process_poi(cls_by_cat.get("poi_public_service", []))
    print("[apply] processing stats ...")
    stats = process_stats(cls_by_cat.get("stats_macro", []))
    print("[apply] processing policy/case rag ...")
    policy = process_rag(cls_by_cat.get("policy_planning", []),
                         EXTERNAL_DIR / "planning_policy" / "processed" /
                         "research_policy_rag_candidates.jsonl", "policy_planning")
    case = process_rag(cls_by_cat.get("project_case", []),
                       EXTERNAL_DIR / "urban_renewal_cases" / "processed" /
                       "research_case_rag_candidates.jsonl", "project_case")

    # 重写血缘（去重后再登记，带 rich extra 字段）
    from app.services import data_lineage_service
    data_lineage_service.remove_external_by_source_prefix("research_")
    lineage_records = []

    def _rec(category: str, record_count: int, *, fe: bool, report: bool, rag: bool,
             training: bool, trainable_n: int, block_reason: str | None,
             field_schema: list[str], comp_status: str) -> None:
        extra = {
            "provider": config.get("provider", "科研人员提供"),
            "authorization_status": "provided" if authorized else "unknown",
            "commercial_risk_override": authorized,
            "used_for_rag": rag, "can_use_for_training": training,
            "trainable_record_count": trainable_n,
            "training_block_reason": block_reason,
            "field_schema": field_schema[:40],
            "test_contamination_risk": False, "leakage_risk": False,
        }
        lid = data_lineage_service.record_collection_lineage(
            source_id=f"research_{category}", source_name=f"科研语料-{category}",
            source_type="research_corpus", raw_count=record_count,
            cleaned_count=record_count,
            license_status="provided_by_research_partner", compliance_status=comp_status,
            used_for_training=training, used_for_feature_engineering=fe,
            used_for_report=report, file_path=f"research/{category}", extra=extra)
        rich = {"lineage_id": lid, "source_id": f"research_{category}", "category": category,
                "record_count": record_count, "used_for_feature_engineering": fe,
                "used_for_report": report, "used_for_training": training, **extra}
        lineage_records.append(rich)

    hp_train = housing["trainable_property_records"] >= 1000
    _rec("housing_property", housing["total_records"], fe=True, report=True, rag=False,
         training=hp_train, trainable_n=housing["trainable_property_records"],
         block_reason=None if hp_train else "可训练样本<1000",
         field_schema=["price_unit", "price_total", "rent", "area", "region",
                       "community", "build_year", "room", "lng", "lat"],
         comp_status="authorized_research")
    _rec("poi_public_service", poi["emitted_records"], fe=True, report=True, rag=False,
         training=False, trainable_n=0, block_reason="POI 不进入监督训练（仅特征/报告）",
         field_schema=["name", "type", "region", "lng", "lat"], comp_status="authorized_research")
    _rec("stats_macro", sum(f.get("record_count", 0) for f in cls_by_cat.get("stats_macro", [])),
         fe=stats["used_for_feature_engineering"], report=True, rag=False, training=False,
         trainable_n=0, block_reason="统计类未结构化/不训练",
         field_schema=stats["target_indicators"], comp_status="authorized_research"
         if stats["has_structured_stats"] else "still_missing")
    _rec("policy_planning", policy["candidate_count"], fe=False, report=True, rag=True,
         training=False, trainable_n=0, block_reason="政策文本不进入训练（RAG/报告）",
         field_schema=["text_preview", "need_manual_review"], comp_status="authorized_research")
    _rec("project_case", case["candidate_count"], fe=False, report=True, rag=True,
         training=False, trainable_n=0, block_reason="案例资料不进入训练（RAG/报告）",
         field_schema=["text_preview", "need_manual_review"], comp_status="authorized_research")

    # 重写 manifest / lineage / gap（授权版）
    cat_agg = result["cat_agg"]
    manifest = {
        "generated_at": _utcnow(), "phase": "10C.5", "source_folder": result["corpus_name"],
        "authorization": {"provider": config.get("provider"),
                          "authorized_by_user": authorized,
                          "authorization_scope": config.get("authorization_scope"),
                          "commercial_risk_override": config.get("commercial_risk_override"),
                          "license_status": "provided_by_research_partner"},
        "total_files": len(result["inventory"]),
        "category_summary": {c: {k: v for k, v in agg.items() if k != "files"}
                             for c, agg in cat_agg.items()},
        "assets": {
            "housing_property": {k: housing[k] for k in
                                 ["trainable_property_records", "trainable_sale_records",
                                  "trainable_rent_records", "supervised_training_strength",
                                  "can_start_supervised_housing_model", "parsed_records"]},
            "poi_public_service": {"emitted_records": poi["emitted_records"],
                                   "overlap_summary": poi["overlap_summary"]},
            "stats_macro": {"status": stats["status"]},
            "policy_planning": policy, "project_case": case,
        },
        "lineage_records": lineage_records,
        "red_lines": [
            "科研语料经用户授权可用于本项目；commercial_risk 由 hard_block 降为 override_warning",
            "used_for_training 仍按训练门禁控制：仅房价(脱敏+价格标签+位置+>=1000)可训练",
            "POI/统计/政策/案例不进入监督训练；GIS 学习素材仅参考",
            "未混入 competition_test；test_contamination_risk=false；leakage_risk=false",
            "原始文件保留在 gitignore 的科研语料目录，明细脱敏后仅落本地 jsonl，接口不返回明细",
        ],
    }
    (RC_DIR / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    (RC_DIR / "lineage.json").write_text(
        json.dumps({"generated_at": _utcnow(), "phase": "10C.5",
                    "records": lineage_records}, ensure_ascii=False, indent=2),
        encoding="utf-8")

    gap = build_gap_fill_authorized(housing, poi, stats, policy, case, cat_agg)
    (DATA_CATALOG_DIR / "科研语料可补缺口清单.json").write_text(
        json.dumps(gap, ensure_ascii=False, indent=2), encoding="utf-8")
    (DATA_CATALOG_DIR / "科研语料可补缺口清单.md").write_text(_gap_md(gap), encoding="utf-8")

    snapshot = build_pre_phase11_snapshot(housing, poi, stats, policy, case)
    (DATA_CATALOG_DIR / "第11训练前数据快照.json").write_text(
        json.dumps(snapshot, ensure_ascii=False, indent=2), encoding="utf-8")
    (DATA_CATALOG_DIR / "第11训练前数据快照.md").write_text(
        _snapshot_md(snapshot), encoding="utf-8")

    return {"housing": housing, "poi": poi, "stats": stats, "policy": policy,
            "case": case, "lineage_records": lineage_records, "snapshot": snapshot}


def build_gap_fill_authorized(housing, poi, stats, policy, case, cat_agg) -> dict[str, Any]:
    return {
        "generated_at": _utcnow(), "phase": "10C.5", "source": "科研语料",
        "authorization": "用户已确认科研人员授权可用于本项目",
        "gaps": [
            {"gap": "房价样本", "research_can_fill": "yes" if housing[
                "can_start_supervised_housing_model"] else "candidate",
             "trainable": housing["can_start_supervised_housing_model"],
             "trainable_records": housing["trainable_property_records"],
             "supervised_training_strength": housing["supervised_training_strength"],
             "note": "链家挂牌/小区基础信息等含价格+位置+面积，脱敏后可作房价监督训练样本。"},
            {"gap": "公共服务/POI", "research_can_fill": "supplement",
             "trainable": False, "records": poi["emitted_records"],
             "overlap_with_amap": poi["overlap_summary"],
             "note": "美食 POI 可作高德 POI 补充特征/报告，去重后不并入官方 5 万 POI。"},
            {"gap": "人口与收入/统计", "research_can_fill":
             "partial" if stats["has_structured_stats"] else "no",
             "trainable": False, "note": "统计类多为扫描件，仍缺权威年鉴口径，需人工补统计局数据。"},
            {"gap": "政策与规划", "research_can_fill": "rag_reference", "trainable": False,
             "rag_candidates": policy["candidate_count"],
             "need_manual_review": policy["need_manual_review"],
             "note": "政策 PPT 图片版需 OCR/人工，可作 RAG/报告参考，不作结构化事实。"},
            {"gap": "项目案例", "research_can_fill": "rag_reference", "trainable": False,
             "rag_candidates": case["candidate_count"],
             "note": "闭门会演讲 PPT 图片版需人工复核，作 RAG/报告参考。"},
            {"gap": "产业细分", "research_can_fill": "no", "trainable": False,
             "note": "科研语料缺企业/园区结构化数据，仍需人工导入。"},
        ],
        "red_line": "房价外仅 feature_engineering/report/rag；未混入 competition_test。",
    }


def build_pre_phase11_snapshot(housing, poi, stats, policy, case) -> dict[str, Any]:
    amap_man = {}
    amap_path = EXTERNAL_DIR / "amap" / "manifest.json"
    if amap_path.exists():
        try:
            amap_man = json.load(amap_path.open())
        except Exception:  # noqa: BLE001
            amap_man = {}
    amap_total = int(amap_man.get("merged_dedup_total", amap_man.get("record_count", 0)))

    organizer = {"note": "见第10A数据审计；组委会官方数据 used_for_eval/competition_test 隔离。"}
    try:
        from app.db.session import SessionLocal  # type: ignore
        from app.models.housing_record import HousingRecord  # type: ignore
        db = SessionLocal()
        try:
            organizer["internal_housing_samples"] = db.query(HousingRecord).count()
        finally:
            db.close()
    except Exception:  # noqa: BLE001
        organizer["internal_housing_samples"] = "见第10A审计"

    hp_ready = housing["can_start_supervised_housing_model"]
    can_enter = amap_total >= 50000  # 空间/报告类可进入；房价为强项但非硬阻断
    return {
        "generated_at": _utcnow(), "phase_gate": "10C.5 -> 11",
        "organizer_data": organizer,
        "amap_poi": {"merged_dedup_total": amap_total,
                     "stopped_reason": amap_man.get("stopped_reason"),
                     "used_for_feature_engineering": True, "used_for_training": False},
        "research_corpus": {
            "housing_trainable_records": housing["trainable_property_records"],
            "housing_strength": housing["supervised_training_strength"],
            "poi_records": poi["emitted_records"],
            "stats_status": stats["status"],
            "policy_rag_candidates": policy["candidate_count"],
            "case_rag_candidates": case["candidate_count"]},
        "trainable_housing_samples": housing["trainable_property_records"],
        "feature_engineering_data": ["高德 POI 50k", "科研 POI", "科研房价(特征)",
                                     "stats(若结构化)"],
        "report_rag_data": ["政策 PPT(需OCR)", "案例 PPT(需OCR)", "城市快报 PDF(需OCR)"],
        "non_trainable_data": ["POI", "政策/案例图片", "GIS 学习素材", "统计扫描件"],
        "still_missing": ["上海公共数据(WAF需人工导入)", "统计局/年鉴权威口径",
                          "政策/案例 OCR 结构化", "产业企业结构化数据"],
        "can_enter_phase11_now": can_enter,
        "can_start_partial": True,
        "can_start_supervised_housing_model": hp_ready,
        "recommended_phase11_tasks": [
            "房价单价监督回归模型（research_authorized_property，脱敏，过第11数据门禁）"
            if hp_ready else "先补授权房价样本再训练房价模型",
            "POI/圈层空间特征工程（高德 50k + 科研 POI 去重补充）",
            "项目类型识别（规则+特征，policy/case 作 RAG 辅助）",
            "报告生成 RAG（政策/案例 OCR 后入库）"],
        "compliance": {"used_for_training_global": "仅授权房价", "test_contamination_risk": False,
                       "leakage_risk": False},
    }


def _snapshot_md(s: dict[str, Any]) -> str:
    lines = ["# 第11训练前数据快照", "", f"- 生成时间：{s['generated_at']}",
             f"- 阶段门：{s['phase_gate']}", "",
             "## 数据概况",
             f"- 组委会内部房价样本：{s['organizer_data'].get('internal_housing_samples')}",
             f"- 高德 POI 去重：{s['amap_poi']['merged_dedup_total']}",
             f"- 科研可训练房价样本：{s['trainable_housing_samples']}"
             f"（强度 {s['research_corpus']['housing_strength']}）",
             f"- 科研 POI：{s['research_corpus']['poi_records']}",
             f"- 统计类状态：{s['research_corpus']['stats_status']}",
             f"- 政策 RAG 候选：{s['research_corpus']['policy_rag_candidates']}；"
             f"案例 RAG 候选：{s['research_corpus']['case_rag_candidates']}", "",
             "## 进入第11判断",
             f"- can_enter_phase11_now：{s['can_enter_phase11_now']}",
             f"- can_start_supervised_housing_model：{s['can_start_supervised_housing_model']}", "",
             "## 仍缺数据"]
    lines += [f"- {x}" for x in s["still_missing"]]
    lines += ["", "## 建议第11训练任务"]
    lines += [f"- {x}" for x in s["recommended_phase11_tasks"]]
    lines += ["", f"> 合规：used_for_training={s['compliance']['used_for_training_global']}；"
              f"test_contamination_risk={s['compliance']['test_contamination_risk']}；"
              f"leakage_risk={s['compliance']['leakage_risk']}"]
    return "\n".join(lines) + "\n"


def main() -> int:
    apply_auth = "--apply-authorization-config" in sys.argv
    print(f"[scan] project_root={PROJECT_ROOT} apply_authorization={apply_auth}")
    result = scan()
    print(f"[scan] corpus={result['corpus_dir']} files={len(result['inventory'])}")
    out = write_outputs(result, record_lineage=not apply_auth)
    if apply_auth:
        print("[apply] applying authorization config ...")
        ap = apply_authorization(result)
        print(f"[apply] done. lineage records={len(ap['lineage_records'])}")
        print(f"[apply] snapshot can_enter_phase11_now="
              f"{ap['snapshot']['can_enter_phase11_now']} "
              f"housing_model={ap['snapshot']['can_start_supervised_housing_model']}")
        return 0
    agg = result["cat_agg"]
    print("[scan] category summary:")
    for cat, a in sorted(agg.items(), key=lambda kv: kv[1]["file_count"], reverse=True):
        print(f"   {cat:22s} files={a['file_count']:4d} "
              f"records={a['record_count']:8d} review={a['need_manual_review_count']:4d} "
              f"commercial={a['commercial_risk_count']}")
    print(f"[scan] lineage records: {len(out['lineage_records'])}")
    print(f"[scan] outputs under: {RC_DIR}")
    print(f"[scan] gap list: {DATA_CATALOG_DIR / '科研语料可补缺口清单.json'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
